"""Fix Demo 8 in PANEL_DEMO_GUIDE.docx: replace DoS payload with exfiltration."""
from docx import Document
import re

path = r"D:\IoT_Device_Fingerprinting_Framework\PANEL_DEMO_GUIDE.docx"
doc = Document(path)

replacements = [
    # Title hint
    (
        "DoS participation — packet rate 80x normal",
        "data exfiltration — upload 1.68 GB, upload-to-download ratio 120×"
    ),
    # Packet count
    ("\"packet_count\": 2240000", "\"packet_count\": 28000"),
    # byte_count
    ("\"byte_count\": 42000000", "\"byte_count\": 2100000000"),
    # packet_rate (32000 → 400)
    ("\"packet_rate\": 32000", "\"packet_rate\": 400"),
    # byte_rate
    ("\"byte_rate\": 48000000", "\"byte_rate\": 36000000"),
    # syn_ratio
    ("\"syn_ratio\": 0.30", "\"syn_ratio\": 0.005"),
    # upload_bytes
    ("\"upload_bytes\": 28000000", "\"upload_bytes\": 1680000000"),
    # upload_download_ratio
    ("\"upload_download_ratio\": 2.0", "\"upload_download_ratio\": 120.0"),
    # Expected response - confidence
    ("\"confidence\": 0.9612", "\"confidence\": 0.9900"),
    # Expected response - anomaly_score
    ("\"anomaly_score\": 0.9481", "\"anomaly_score\": 0.8060"),
    # Expected response - severity
    ("\"severity\": \"Critical\"", "\"severity\": \"High\""),
    # Say to panel text
    (
        "The /analyze endpoint does everything in one call — it first fingerprints the device as a smart camera with 96% confidence, then immediately scores it for anomalies. The camera is sending SYN floods at 32,000 packets per second — 80 times its normal rate. This is a DoS participation attack. Both results are returned together in a single response.",
        "The /analyze endpoint does everything in one call. It first fingerprints the device automatically as a smart camera with 99% confidence — no device type was provided by the caller. Then it immediately scores the same traffic for anomalies. The upload-to-download ratio is 120, whereas a camera’s normal ratio is around 2. The camera is sending 1.68 GB out but only receiving 14 MB — a classic data exfiltration signature. The system raises a High-severity alert at score 0.806 without any manual configuration."
    ),
]

changes = 0
for para in doc.paragraphs:
    for old, new in replacements:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    changes += 1
            # fallback: rebuild paragraph text across runs
            if old in para.text:
                full = para.text
                if old in full:
                    # clear all runs, put in first run
                    para.runs[0].text = full.replace(old, new)
                    for r in para.runs[1:]:
                        r.text = ""
                    changes += 1

# Also check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in replacements:
                    if old in para.text:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                changes += 1

doc.save(path)
print(f"Done — {changes} replacements made in {path}")
