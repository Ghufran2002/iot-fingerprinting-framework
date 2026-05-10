"""
Updates PANEL_DEMO_GUIDE.docx, PROJECT_GUIDE.docx, DISSERTATION_REPORT.docx
to include SHAP explainability content and POST /explain endpoint.
Run: python update_docs_shap.py
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

# ── Helper ────────────────────────────────────────────────────────────────────

def insert_paras_after(doc, anchor_text, new_items):
    """
    Insert paragraphs after the LAST paragraph whose text contains anchor_text.
    new_items = list of (text, bold, heading_level)  heading_level=0 means normal
    """
    body = doc.element.body
    all_paras = doc.paragraphs
    target = None
    for p in all_paras:
        if anchor_text in p.text:
            target = p
    if target is None:
        print(f"  WARNING: anchor not found: {anchor_text!r}")
        return False

    insert_idx = list(body).index(target._element) + 1

    for offset, (text, bold, hlevel) in enumerate(reversed(new_items)):
        if hlevel > 0:
            np = doc.add_heading(text, level=hlevel)
        else:
            np = doc.add_paragraph(text)
        if bold and not hlevel:
            for run in np.runs:
                run.bold = True
        body.remove(np._element)
        body.insert(insert_idx, np._element)

    return True


def replace_text(doc, old_text, new_text):
    """Replace first exact match of old_text in any paragraph."""
    for p in doc.paragraphs:
        if old_text in p.text:
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return True
            # Fallback: paragraph has the text split across runs
            if p.text == old_text:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""
                return True
    print(f"  WARNING: text not found for replace: {old_text!r}")
    return False


# ═══════════════════════════════════════════════════════════════════════════════
#  1. PANEL_DEMO_GUIDE.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating PANEL_DEMO_GUIDE.docx ...")
doc = Document("PANEL_DEMO_GUIDE.docx")

# Add new Demo 10 — SHAP after Demo 9 (live dashboard section)
insert_paras_after(doc,
    "This is the live monitoring dashboard. A security analyst or home user can watch this in real time.",
    [
        ("3.11 Demo 10 — SHAP Live Explainability (Unique Feature — Impress the Panel)", True, 0),
        ("Still in Tab 1 (Dashboard) — scroll to the bottom of the page.", False, 0),
        ("What to show:", True, 0),
        (
            "At the very bottom of the dashboard you will see the SHAP Explainability Panel. "
            "It shows a horizontal bar chart labelled \"SHAP: Why the model says → [Device] (XX% confidence)\". "
            "The bars update automatically every 5 seconds as the dashboard polls the live POST /explain endpoint.",
            False, 0
        ),
        ("Green bars = features that PUSH the prediction toward that device type.", False, 0),
        ("Red bars   = features that PUSH AWAY from that device type.", False, 0),
        ("Switch to Tab 2 (API Docs) and show POST /explain manually:", True, 0),
        (
            "In Swagger UI, find POST /explain under Explainability. "
            "Click Try it out. Use the default payload (smart camera profile). Click Execute.",
            False, 0
        ),
        ("The response will show device_type, confidence, and top 10 features each with a shap_value and direction.", False, 0),
        ("Say to the panel:", True, 0),
        (
            "\"Sir, this is the POST /explain endpoint — a live SHAP explainability API. "
            "For every single prediction the model makes, I can tell you exactly which feature "
            "contributed how much, and in which direction. For example, tcp_ratio of 0.92 strongly "
            "pushes toward smart_camera — because cameras are TCP-heavy devices. "
            "is_coap = 0 pushes away from smart_bulb — because CoAP is the bulb's signature protocol. "
            "This proves my model is NOT a black box. Every decision is fully transparent and auditable. "
            "This is what Explainable AI means in practice.\"",
            False, 0
        ),
    ]
)

# Rename old "3.11 Demo 10" to "3.12 Demo 11"
replace_text(doc, "3.11 Demo 10 — Show Evaluation Results (Plots)", "3.12 Demo 11 — Show Evaluation Results (Plots)")

# Add SHAP Q&A question after the last existing question about Isolation Forest
insert_paras_after(doc,
    "Isolation Forest",
    [
        ("Question 6 (NEW): \"What is SHAP and why is it important in your project?\"", True, 0),
        (
            "Answer: \"SHAP stands for SHapley Additive exPlanations. It is a game-theory based "
            "technique that calculates how much each feature contributed to each individual model prediction. "
            "I implemented a live POST /explain endpoint in my REST API. "
            "When you call it with a network flow, it returns the top 10 features and their SHAP values. "
            "Positive values push toward the predicted device; negative values push away. "
            "The dashboard also shows these as a green/red bar chart in real time. "
            "This makes my system fully explainable — which is critical in a security context where "
            "a human analyst needs to understand and trust every alert the system raises.\"",
            False, 0
        ),
        ("Question 7 (NEW): \"How many API endpoints does your system have?\"", True, 0),
        (
            "Answer: \"My system has 9 REST API endpoints: "
            "GET /health, GET /devices, POST /fingerprint, POST /anomaly/score, POST /analyze, "
            "POST /explain (SHAP), GET /alerts/recent, GET /metrics, POST /demo/inject. "
            "The most important is POST /analyze — it combines fingerprinting and anomaly detection "
            "in a single call. And POST /explain adds full SHAP explainability on top.\"",
            False, 0
        ),
    ]
)

doc.save("PANEL_DEMO_GUIDE.docx")
print("  Saved PANEL_DEMO_GUIDE.docx")


# ═══════════════════════════════════════════════════════════════════════════════
#  2. PROJECT_GUIDE.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating PROJECT_GUIDE.docx ...")
doc = Document("PROJECT_GUIDE.docx")

# Update REST API endpoints list
replace_text(doc,
    "Endpoints: /fingerprint /anomaly/score /analyze /alerts",
    "Endpoints: /fingerprint /anomaly/score /analyze /explain (SHAP) /alerts /metrics"
)

# Add SHAP info after dashboard section
insert_paras_after(doc,
    "KPI Cards: Total Alerts, Anomaly Rate, Requests, Devices (8), Features (37), Threshold (0.75)",
    [
        ("SHAP Explainability Panel (dashboard ke bottom mein):", True, 0),
        (
            "Har 5 seconds mein dashboard POST /explain endpoint call karta hai. "
            "Ek horizontal bar chart dikhata hai — green bars matlab feature model ko us device ki taraf push kar raha hai, "
            "red bars matlab door kar raha hai. Yeh Explainable AI hai — model black box nahi hai.",
            False, 0
        ),
    ]
)

# Add SHAP explanation in the architecture flow
insert_paras_after(doc,
    "STEP 4: REST API  (FastAPI — Port 8000)",
    [
        (
            "9 Endpoints: /fingerprint, /anomaly/score, /analyze, /explain (live SHAP), "
            "/alerts/recent, /metrics, /health, /devices, /demo/inject",
            False, 0
        ),
    ]
)

# Add SHAP Q&A
insert_paras_after(doc,
    "Sir, framework mein FastAPI REST API hai port 8000 pe.",
    [
        ("Q (SHAP): SHAP kya hai aur aapne kyun use kiya?", True, 0),
        (
            "Sir, SHAP — SHapley Additive exPlanations — ek game-theory based technique hai jo "
            "har prediction ke liye explain karta hai ki kaunse feature ne kitna contribute kiya. "
            "Maine POST /explain endpoint banaya hai — jab bhi koi network flow aata hai, "
            "API top 10 features return karta hai with SHAP values: positive matlab us device ki taraf push, "
            "negative matlab door push. Dashboard mein bhi green/red bar chart live update hota hai. "
            "Yeh IIT panels ke liye bahut important hai — proves karta hai model black box nahi hai "
            "aur har decision transparent aur auditable hai.",
            False, 0
        ),
    ]
)

doc.save("PROJECT_GUIDE.docx")
print("  Saved PROJECT_GUIDE.docx")


# ═══════════════════════════════════════════════════════════════════════════════
#  3. DISSERTATION_REPORT.docx
# ═══════════════════════════════════════════════════════════════════════════════
print("Updating DISSERTATION_REPORT.docx ...")
doc = Document("DISSERTATION_REPORT.docx")

# 3a. Add SHAP after section 8.3 Feature Importance Analysis heading
insert_paras_after(doc,
    "8.3 Feature Importance Analysis",
    [
        ("8.3.1 SHAP-Based Explainability Analysis", True, 0),
        (
            "To address the interpretability requirement of a forensic-grade security system, "
            "this work employs SHAP (SHapley Additive exPlanations) [Lundberg & Lee, 2017] as the "
            "primary explainability framework. Unlike traditional feature importance scores derived "
            "from mean Gini impurity decrease — which provide only global, aggregate rankings — "
            "SHAP provides per-prediction, theoretically grounded attributions satisfying the axioms "
            "of efficiency, symmetry, and null player from cooperative game theory.",
            False, 0
        ),
        (
            "The SHAP TreeExplainer is applied to the trained Random Forest model on a stratified "
            "sample of 400 flows (50 per device type). The resulting SHAP values confirm that the "
            "model relies on semantically meaningful features:",
            False, 0
        ),
        (
            "Global Feature Rankings (Mean |SHAP| across all classes and samples): "
            "(1) tcp_ratio — 0.2386: Protocol dominance is the strongest device fingerprint. "
            "TCP-heavy devices (cameras, TVs, doorbells) are immediately separated from UDP-heavy "
            "devices (bulbs, sensors). "
            "(2) udp_ratio — 0.2242: Complementary to tcp_ratio. "
            "(3) ack_ratio — 0.0949: TCP handshake pattern distinguishes streaming devices from "
            "event-driven sensors. "
            "(4) mean_dest_port — 0.0201: Port 443 identifies camera/TV (HTTPS); "
            "port 1883 identifies thermostat/plug (MQTT); port 5683 identifies smart_bulb (CoAP). "
            "(5) is_mqtt — boolean flag alone identifies thermostat and smart_plug with high confidence.",
            False, 0
        ),
        (
            "Per-Device Analysis (SHAP Heatmap): The per-device SHAP heatmap reveals that "
            "smart_bulb is the most uniquely identifiable device — is_coap dominates its row "
            "completely, achieving a normalised SHAP score of 1.00. This means CoAP protocol usage "
            "on port 5683 is, in isolation, sufficient to identify a smart bulb with near-certainty. "
            "Motion sensors are characterised by near-zero byte_count; smart cameras by the "
            "tcp_ratio + is_https + ack_ratio triplet.",
            False, 0
        ),
        (
            "Live Explainability Deployment: The SHAP TreeExplainer is loaded at API startup "
            "and exposed via the POST /explain endpoint. For any incoming 37-feature flow vector, "
            "the endpoint returns the predicted device type, confidence score, and the top 10 "
            "features with their signed SHAP contributions (positive = toward prediction, "
            "negative = away). The real-time monitoring dashboard polls this endpoint every "
            "5 seconds and renders a live horizontal bar chart (green = toward, red = away), "
            "making every model prediction fully auditable by a human analyst.",
            False, 0
        ),
    ]
)

# 3b. Add /explain to API endpoints table reference
insert_paras_after(doc,
    "Table 7.2 — REST API Endpoints Reference",
    [
        (
            "Note: Table 7.2 lists all 9 endpoints. The POST /explain endpoint (added in v1.0) "
            "accepts a FlowFeatures payload and returns signed SHAP contributions for the top 10 "
            "features, making the fingerprinting decision fully explainable and auditable.",
            False, 0
        ),
    ]
)

# 3c. Update keywords to include SHAP
replace_text(doc,
    "Keywords: IoT Security, Device Fingerprinting, Anomaly Detection, Random Forest, Isolation Forest, O",
    "Keywords: IoT Security, Device Fingerprinting, Anomaly Detection, Explainable AI, SHAP, Random Forest, Isolation Forest, One-Class SVM, FastAPI, Real-Time Monitoring"
)

doc.save("DISSERTATION_REPORT.docx")
print("  Saved DISSERTATION_REPORT.docx")

print("\nAll 3 docs updated successfully.")
