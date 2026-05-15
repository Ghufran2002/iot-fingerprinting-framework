"""Fix outdated references in DISSERTATION_REPORT_FINAL.docx"""
from docx import Document
import copy

REPLACEMENTS = [
    # Plotly Dash -> Plotly.js
    ("Plotly Dash Dashboard",       "Live Monitoring Dashboard (Plotly.js)"),
    ("Plotly Dash dashboard",        "live Plotly.js dashboard"),
    ("Plotly Dash live dashboard",   "live Plotly.js monitoring dashboard"),
    ("Plotly Dash",                  "Plotly.js"),
    ("plotly dash",                  "Plotly.js"),
    # Ports
    ("port 8000 and the Plotly Dash dashboard (port 8050)", "port 7860"),
    ("port 8050",                    "port 7860"),
    ("port 8000",                    "port 7860"),
    ("(port 8000)",                  "(port 7860)"),
    ("FastAPI server (port 8000) and the Plotly Dash dashboard (port 8050)",
     "FastAPI server on port 7860 (API + Dashboard on single port)"),
    # Source files
    ("src/dashboard/app.py",         "src/api/main.py"),
    # Confidence threshold
    ("CONFIDENCE_THRESHOLD = 0.75",  "CONFIDENCE_THRESHOLD = 0.60"),
    # Scaling formula
    ("self._if_hi = float(np.percentile(if_train, 95) * 2.0)",
     "self._if_hi = if_p95 + max(if_p95 - if_p5, 0.05)"),
    # Dashboard URL references
    ("Plotly Dash dashboard at https://mdghufran-iot-fingerprinting.hf.space/dashboard/",
     "live monitoring dashboard at https://mdghufran-iot-fingerprinting.hf.space/dashboard/"),
    # Figure captions
    ("Figure 7.4: Plotly Dash Dashboard",   "Figure 7.4: Live Monitoring Dashboard"),
    ("Figure 7.4 — Plotly Dash Dashboard",  "Figure 7.4 — Live Monitoring Dashboard (Plotly.js)"),
]

def replace_in_runs(para, old, new):
    """Replace text in paragraph runs. Handles single-run matches."""
    full = para.text
    if old not in full:
        return 0
    count = 0
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    return count

def fix_paragraph(para):
    total = 0
    for old, new in REPLACEMENTS:
        total += replace_in_runs(para, old, new)
    return total

def fix_doc(path):
    doc = Document(path)
    total = 0

    # Body paragraphs
    for para in doc.paragraphs:
        total += fix_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += fix_paragraph(para)

    doc.save(path)
    print(f"Done — {total} replacements made in {path}")

fix_doc('DISSERTATION_REPORT_FINAL.docx')
