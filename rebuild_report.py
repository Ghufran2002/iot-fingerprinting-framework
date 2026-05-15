"""
Rebuild DISSERTATION_REPORT.docx as a professional university-grade document.
- Converts all bullet-point sections to flowing academic prose
- Replaces ASCII architecture with proper diagrams
- Inserts all generated plots
- Applies consistent academic formatting
"""
import copy, io, shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

SRC  = 'DISSERTATION_REPORT.docx'
DEST = 'DISSERTATION_REPORT_FINAL.docx'

# ── helpers ──────────────────────────────────────────────────────────
def set_para_spacing(para, before=0, after=6, line=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line

def body_para(doc, text, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.style = doc.styles['Body Text']
    set_para_spacing(p, before=0, after=after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    return p

def first_para(doc, text):
    return body_para(doc, text)

def add_figure(doc, img_path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=6, after=2)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap, before=2, after=12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    set_para_spacing(p, before=24, after=12)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    set_para_spacing(p, before=16, after=8)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    set_para_spacing(p, before=12, after=6)
    return p

def page_break(doc):
    doc.add_page_break()

# ── open source and extract embedded images ───────────────────────────
src_doc = Document(SRC)

# Pull existing embedded images (screenshots) by name
embedded = {}
for rel in src_doc.part.rels.values():
    if 'image' in rel.reltype:
        name = rel.target_ref.split('/')[-1]
        embedded[name] = rel.target_part.blob

# Screenshots in order they appear: image1..image6
screenshot_blobs = [embedded.get(f'image{i}.png') for i in range(1, 7)]

# ── open source to mine original table XML ───────────────────────────
src_tables = src_doc.tables   # we'll copy tables where needed

def clone_table(src_tbl, dst_doc):
    """Deep-copy a table from source into destination document."""
    tbl_copy = copy.deepcopy(src_tbl._tbl)
    dst_doc.element.body.append(tbl_copy)

# Map table index in source doc to a label
# (inspect source manually — tables appear in chapter order)
# We will copy them in sequence as we write each section.
TABLE_IDX = 0  # counter reset per use in write_section

def next_table(dst_doc):
    global TABLE_IDX
    if TABLE_IDX < len(src_tables):
        clone_table(src_tables[TABLE_IDX], dst_doc)
        TABLE_IDX += 1

def add_screenshot(dst_doc, blob, caption, width=5.8):
    if blob is None:
        return
    p = dst_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=6, after=2)
    run = p.add_run()
    from docx.shared import Inches
    import io
    run.add_picture(io.BytesIO(blob), width=Inches(width))
    cap = dst_doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap, before=2, after=12)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ════════════════════════════════════════════════════════════════════
# BUILD NEW DOCUMENT
# ════════════════════════════════════════════════════════════════════
doc = Document()

# ── page margins ──────────────────────────────────────────────────
from docx.oxml.ns import qn
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ── default font ──────────────────────────────────────────────────
from docx.shared import Pt
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ensure Body Text style exists and is formatted
for sname in ['Body Text', 'First Paragraph']:
    if sname in doc.styles:
        s = doc.styles[sname]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════
def title_page(doc):
    for _ in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('DISSERTATION REPORT')
    r.bold = True; r.font.size = Pt(18)
    set_para_spacing(t, after=18)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        'DESIGN OF A FRAMEWORK FOR IoT DEVICE FINGERPRINTING AND ANOMALY\n'
        'DETECTION FOR SMART HOME USING MACHINE LEARNING')
    r2.bold = True; r2.font.size = Pt(14)
    set_para_spacing(t2, before=6, after=24)

    for line in [
        'Submitted in partial fulfilment of the requirements for the degree of',
        'Master of Technology in Cyber Forensics',
        'Session: 2024–2026 (4th Semester)',
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=3)

    doc.add_paragraph()
    for line in ['Submitted by:', 'Md Ghufran Alam', 'Roll Number: NDU202400038']:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=3)
        if line == 'Submitted by:':
            p.runs[0].bold = True

    doc.add_paragraph()
    for line in [
        'Under the Supervision of:',
        'Dr. Syed Nisar Hussain Bhukhari',
        'NDU Coordinator, Department of CSE',
        'NIELIT Srinagar',
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=3)
        if line == 'Under the Supervision of:':
            p.runs[0].bold = True

    for _ in range(4):
        doc.add_paragraph()

    footer = doc.add_paragraph(
        'NATIONAL INSTITUTE OF ELECTRONICS AND INFORMATION TECHNOLOGY (NIELIT)\n'
        'SRINAGAR, JAMMU & KASHMIR\nMAY 2026')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].bold = True
    set_para_spacing(footer, after=0)

title_page(doc)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# DECLARATION
# ════════════════════════════════════════════════════════════════════
h2(doc, 'DECLARATION')
first_para(doc,
    'I, Md Ghufran Alam, Roll Number NDU202400038, a student of M.Tech in Cyber Forensics '
    '(Session 2024–2026) at the National Institute of Electronics and Information Technology '
    '(NIELIT) Srinagar, hereby declare that the dissertation entitled “Design of a Framework '
    'for IoT Device Fingerprinting and Anomaly Detection for Smart Home Using Machine Learning” '
    'submitted in partial fulfilment of the requirements for the degree of Master of Technology is a '
    'record of original work done by me under the guidance of my supervisor.')
body_para(doc,
    'I further declare that this dissertation or any part thereof has not been submitted for the award '
    'of any other degree or diploma in this Institute or any other Institution/University. The '
    'information or data collected from other sources has been duly acknowledged in the text.')
body_para(doc, 'Place: Srinagar, J&K\t\t\t\t\t\t\tDate: May 2026', after=24)
body_para(doc, 'Md Ghufran Alam\n(NDU202400038)')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CERTIFICATE
# ════════════════════════════════════════════════════════════════════
h2(doc, 'CERTIFICATE')
first_para(doc,
    'This is to certify that the dissertation entitled “Design of a Framework for IoT Device '
    'Fingerprinting and Anomaly Detection for Smart Home Using Machine Learning” submitted by '
    'Md Ghufran Alam (Roll No. NDU202400038) in partial fulfilment of the requirements for the award '
    'of the degree of Master of Technology in Cyber Forensics from the National Institute of '
    'Electronics and Information Technology (NIELIT) Srinagar is a record of bonafide research work '
    'carried out by him under my supervision.')
body_para(doc,
    'The results embodied in this dissertation have not been submitted for the award of any other '
    'degree or diploma.')
body_para(doc, 'Place: Srinagar, J&K\t\t\t\t\t\t\tDate: May 2026', after=36)
body_para(doc,
    'Dr. Syed Nisar Hussain Bhukhari\nNDU Coordinator, Department of CSE\nNIELIT Srinagar',
    after=36)
body_para(doc, 'Head of Department / Director\nNIELIT Srinagar')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ════════════════════════════════════════════════════════════════════
h2(doc, 'ACKNOWLEDGEMENTS')
for para_text in [
    ('I would like to express my sincere gratitude to everyone who supported and guided me throughout '
     'this research work.'),
    ('First and foremost, I thank my project supervisor, Dr. Syed Nisar Hussain Bhukhari (NDU '
     'Coordinator, Department of CSE, NIELIT Srinagar), for providing invaluable guidance, '
     'encouragement, and constructive feedback at every stage of this dissertation. Their expertise in '
     'cyber security and machine learning has been instrumental in shaping the scope and direction of '
     'this work.'),
    ('I am grateful to the Director and faculty of NIELIT Srinagar for providing the academic '
     'environment and resources necessary to carry out this research. Special thanks are due to the '
     'Department of Cyber Forensics for the technical support and domain knowledge imparted throughout '
     'the programme.'),
    ('I would also like to acknowledge the open-source community behind scikit-learn, FastAPI, Plotly '
     'Dash, and the N-BaIoT dataset researchers at Ben-Gurion University for making their datasets '
     'publicly available, which greatly enriched the validation of this framework.'),
    ('Finally, I owe a debt of gratitude to my family for their patience, moral support, and '
     'unwavering encouragement throughout the duration of this programme.'),
    ('Md Ghufran Alam\nNIELIT Srinagar, May 2026'),
]:
    body_para(doc, para_text)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════
h2(doc, 'ABSTRACT')
for para_text in [
    ('The proliferation of Internet of Things (IoT) devices in smart home environments has introduced '
     'unprecedented security challenges. Unlike conventional computing systems, IoT devices are '
     'resource-constrained, heterogeneous in design, and rarely receive timely security patches, making '
     'them prime targets for botnets, data exfiltration campaigns, and denial-of-service attacks. '
     'Existing network-level security solutions, designed for traditional IT infrastructure, are largely '
     'ineffective against the unique traffic patterns exhibited by IoT endpoints.'),
    ('This dissertation presents the design and implementation of a modular, end-to-end framework for '
     'IoT device fingerprinting and anomaly detection in smart home networks using machine learning. The '
     'framework addresses two tightly coupled problems: (i) passive device identification '
     '(fingerprinting) from network flow features, and (ii) per-device behavioural anomaly detection to '
     'flag deviations from a learned baseline.'),
    ('The fingerprinting subsystem employs a Random Forest classifier (100 estimators, max_depth = 15) '
     'alongside Gradient Boosting, Support Vector Machine, and a soft-voting Ensemble, trained on 37 '
     'statistical and protocol-level features extracted from network flows across 8 IoT device '
     'categories. The anomaly detection subsystem deploys a per-device ensemble of Isolation Forest (IF) '
     'and One-Class SVM (OC-SVM), with ensemble weights of 60 % IF and 40 % OC-SVM, calibrated '
     'to an alert threshold of 0.75.'),
    ('A synthetic dataset of 1,600 labelled flows (200 per device, 5 % anomaly rate) was constructed '
     'for training and evaluation, supplemented by real traffic samples from the N-BaIoT dataset. Class '
     'imbalance is addressed using SMOTE oversampling, and features are normalised with RobustScaler. '
     'The trained models are exposed via a FastAPI REST service (port 8000) and a Plotly Dash '
     'monitoring dashboard (port 8050) for real-time operational use.'),
    ('Experimental results demonstrate that the Random Forest fingerprinter achieves test accuracy of '
     '100 % with a macro ROC-AUC of 1.0000 across all device classes. The per-device anomaly '
     'detectors consistently separate anomalous from normal traffic, with mean anomaly scores for attack '
     'flows exceeding 0.85 against a threshold of 0.75. The framework is entirely open-source, '
     'reproducible, and deployable on commodity hardware with 8 GB RAM.'),
]:
    body_para(doc, para_text)

kw = doc.add_paragraph()
kw.style = doc.styles['Body Text']
set_para_spacing(kw, before=6, after=0)
r = kw.add_run('Keywords: ')
r.bold = True
kw.add_run(
    'IoT Security, Device Fingerprinting, Anomaly Detection, Explainable AI, SHAP, '
    'Random Forest, Isolation Forest, One-Class SVM, Smart Home, Network Forensics, SMOTE, FastAPI')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (carried from source)
# ════════════════════════════════════════════════════════════════════
h2(doc, 'TABLE OF CONTENTS')
toc_entries = [
    ('Chapter 1: Introduction', ''),
    ('    1.1  Background and Motivation', ''),
    ('    1.2  Problem Statement', ''),
    ('    1.3  Objectives', ''),
    ('    1.4  Scope and Delimitations', ''),
    ('    1.5  Dissertation Organisation', ''),
    ('Chapter 2: Literature Review', ''),
    ('    2.1  IoT Security: An Overview', ''),
    ('    2.2  IoT Device Fingerprinting Techniques', ''),
    ('    2.3  Anomaly Detection in IoT Networks', ''),
    ('    2.4  Machine Learning for Network Intrusion Detection', ''),
    ('    2.5  Relevant Datasets', ''),
    ('    2.6  Research Gaps and Positioning', ''),
    ('Chapter 3: System Architecture and Design', ''),
    ('    3.1  Design Principles', ''),
    ('    3.2  High-Level Architecture', ''),
    ('    3.3  Module Descriptions', ''),
    ('    3.4  Data Flow', ''),
    ('Chapter 4: Feature Engineering', ''),
    ('    4.1  Feature Selection Rationale', ''),
    ('    4.2  Feature Categories and Definitions', ''),
    ('    4.3  Per-Device Traffic Characteristics', ''),
    ('Chapter 5: Dataset Construction and Preprocessing', ''),
    ('    5.1  Synthetic Dataset Generation', ''),
    ('    5.2  Real-World N-BaIoT Data Integration', ''),
    ('    5.3  Anomaly Injection', ''),
    ('    5.4  Preprocessing Pipeline', ''),
    ('Chapter 6: Machine Learning Models', ''),
    ('    6.1  Device Fingerprinting Models', ''),
    ('    6.2  Anomaly Detection Ensemble', ''),
    ('    6.3  Model Persistence and Deployment', ''),
    ('Chapter 7: System Implementation', ''),
    ('    7.1  Technology Stack', ''),
    ('    7.2  Project Structure', ''),
    ('    7.3  Training Pipeline', ''),
    ('    7.4  REST API (FastAPI)', ''),
    ('    7.5  Monitoring Dashboard (Plotly Dash)', ''),
    ('    7.6  Live System Demonstration', ''),
    ('Chapter 8: Results and Evaluation', ''),
    ('    8.1  Fingerprinting Performance', ''),
    ('    8.2  Anomaly Detection Performance', ''),
    ('    8.3  Feature Importance Analysis', ''),
    ('    8.4  Comparative Analysis', ''),
    ('Chapter 9: Conclusion and Future Work', ''),
    ('    9.1  Summary of Contributions', ''),
    ('    9.2  Limitations', ''),
    ('    9.3  Future Directions', ''),
    ('References', ''),
    ('Appendix A — Feature Definitions', ''),
    ('Appendix B — API Endpoint Reference', ''),
    ('Appendix C — System Requirements and Setup', ''),
]
for entry, _ in toc_entries:
    p = doc.add_paragraph(entry)
    p.style = doc.styles['Normal']
    set_para_spacing(p, before=0, after=2)
    p.runs[0].font.size = Pt(11)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# LIST OF FIGURES / TABLES / ABBREVIATIONS  (from source)
# ════════════════════════════════════════════════════════════════════
h2(doc, 'LIST OF FIGURES')
figs = [
    'Figure 3.1 — High-Level System Architecture',
    'Figure 3.2 — End-to-End Data Flow Diagram',
    'Figure 4.1 — Feature Category Distribution (37 Features)',
    'Figure 5.1 — Dataset Device Distribution',
    'Figure 5.2 — SMOTE Oversampling Effect on Class Balance',
    'Figure 6.1 — Random Forest Decision Ensemble Architecture',
    'Figure 6.2 — Per-Device Anomaly Detector Architecture',
    'Figure 7.1 — FastAPI Swagger UI — Endpoint Overview (upper section)',
    'Figure 7.2 — FastAPI Swagger UI — POST /analyze, /explain and Schemas',
    'Figure 7.3 — GET /health JSON Response — HTTP 200, 37 ms, uptime 7551 s',
    'Figure 7.4 — Plotly Dash Dashboard — KPI Strip and Anomaly Timeline',
    'Figure 7.5 — Dashboard — Alert Severity, Score Gauge and Alerts Table',
    'Figure 7.6 — SHAP Live Explainability Panel — Top-10 Feature Contributions',
    'Figure 8.1 — Confusion Matrix — Random Forest',
    'Figure 8.2 — Confusion Matrix — Gradient Boosting',
    'Figure 8.3 — Confusion Matrix — SVM',
    'Figure 8.4 — Confusion Matrix — Voting Ensemble',
    'Figure 8.5 — ROC Curves (One-vs-Rest) — Random Forest',
    'Figure 8.6 — Top-15 Feature Importances (Gini) — Random Forest',
    'Figure 8.7/8.8 — Model Comparison — Test Accuracy and Macro ROC-AUC',
    'Figure 8.9 — Per-Device Anomaly Score Distribution',
]
for f in figs:
    p = doc.add_paragraph(f); p.style = doc.styles['Normal']
    set_para_spacing(p, before=0, after=2); p.runs[0].font.size = Pt(11)

doc.add_paragraph()
h2(doc, 'LIST OF TABLES')
tbls = [
    'Table 3.1 — Module Summary',
    'Table 4.1 — Feature Categories and Definitions',
    'Table 4.2 — Per-Device Protocol Signatures',
    'Table 5.1 — Dataset Composition Summary',
    'Table 5.2 — Anomaly Type Definitions and Injection Multipliers',
    'Table 6.1 — Fingerprinting Model Hyperparameters',
    'Table 6.2 — Anomaly Detector Hyperparameters',
    'Table 7.1 — Technology Stack Summary',
    'Table 7.2 — REST API Endpoints Reference',
    'Table 8.1 — Fingerprinting Classification Report (Macro Averages)',
    'Table 8.2 — Per-Model Test Accuracy and ROC-AUC',
    'Table 8.3 — Per-Device Anomaly Detection Metrics',
    'Table 8.4 — Comparison with Related Works',
]
for t in tbls:
    p = doc.add_paragraph(t); p.style = doc.styles['Normal']
    set_para_spacing(p, before=0, after=2); p.runs[0].font.size = Pt(11)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ════════════════════════════════════════════════════════════════════
h2(doc, 'LIST OF ABBREVIATIONS')
abbrevs = [
    ('AI',             'Artificial Intelligence'),
    ('API',            'Application Programming Interface'),
    ('AUC',            'Area Under the Curve'),
    ('CoAP',           'Constrained Application Protocol'),
    ('CORS',           'Cross-Origin Resource Sharing'),
    ('DPI',            'Deep Packet Inspection'),
    ('DoS',            'Denial of Service'),
    ('HTTP/HTTPS',     'Hypertext Transfer Protocol (Secure)'),
    ('IF',             'Isolation Forest'),
    ('IoT',            'Internet of Things'),
    ('IP',             'Internet Protocol'),
    ('JSON',           'JavaScript Object Notation'),
    ('KPI',            'Key Performance Indicator'),
    ('ML',             'Machine Learning'),
    ('MQTT',           'Message Queuing Telemetry Transport'),
    ('NTP',            'Network Time Protocol'),
    ('OC-SVM',         'One-Class Support Vector Machine'),
    ('REST',           'Representational State Transfer'),
    ('RF',             'Random Forest'),
    ('ROC',            'Receiver Operating Characteristic'),
    ('RobustScaler',   'Robust feature normalisation using interquartile range'),
    ('SHAP',           'SHapley Additive exPlanations'),
    ('SMOTE',          'Synthetic Minority Over-sampling Technique'),
    ('SVM',            'Support Vector Machine'),
    ('TCP',            'Transmission Control Protocol'),
    ('UDP',            'User Datagram Protocol'),
    ('XAI',            'Explainable Artificial Intelligence'),
]
for abbr, defn in abbrevs:
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    set_para_spacing(p, before=0, after=2)
    r = p.add_run(f'{abbr:<16}'); r.bold = True; r.font.name = 'Courier New'; r.font.size = Pt(11)
    p.add_run(defn).font.size = Pt(11)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 1: INTRODUCTION')
h2(doc, '1.1  Background and Motivation')
for t in [
    ('The Internet of Things (IoT) has fundamentally transformed the concept of the modern home. From '
     'smart cameras that stream live video to cloud servers, to thermostats that adjust temperature '
     'based on occupancy patterns, connected devices now permeate residential environments at an '
     'accelerating rate. Ericsson (2023) estimates that the global number of IoT connections will '
     'exceed 34 billion by 2028, with smart home devices constituting a substantial and growing '
     'fraction of that total.'),
    ('This explosive growth, however, has outpaced the development of security safeguards. IoT devices '
     'are architecturally distinct from conventional computing systems: they operate on stripped-down '
     'firmware, lack persistent storage for security patches, and are manufactured by vendors for whom '
     'security is rarely a primary design consideration. The resulting security posture is '
     'systematically weak, creating an attack surface that is simultaneously large, heterogeneous, and '
     'difficult to monitor.'),
    ('The consequences are well-documented and severe. The Mirai botnet (2016) co-opted over 600,000 '
     'IoT devices — primarily home routers, IP cameras, and digital video recorders — to '
     'launch distributed denial-of-service attacks that disrupted major internet services globally '
     '(Kolias et al., 2017). Subsequent variants of Mirai and related botnets have demonstrated that '
     'the attack vector remains unresolved.'),
    ('A foundational requirement for defending smart home networks is knowing precisely what devices '
     'are present on the network and whether each device is behaving as expected. Without device '
     'identity, administrators cannot enforce per-device access control policies or detect anomalous '
     'behaviour. Traditional network security tools are ill-suited to this task: IoT devices do not '
     'support agent-based monitoring, do not respond reliably to active probing, and exhibit traffic '
     'patterns that differ fundamentally from those of workstations or servers.'),
    ('This dissertation proposes a passive, non-intrusive solution: a machine learning-based framework '
     'that fingerprints IoT devices solely from statistical and protocol-level features derived from '
     'network flows, and simultaneously monitors each identified device for behavioural anomalies '
     'indicative of compromise. The framework requires no modification to device firmware, no '
     'installation of agents, and no access to encrypted packet payloads.'),
]:
    body_para(doc, t)

h2(doc, '1.2  Problem Statement')
first_para(doc,
    'Despite significant research effort, practical IoT security in the smart home context remains '
    'unsolved for two primary reasons. First, identity ambiguity means that network administrators '
    'cannot reliably identify what IoT device type is responsible for a given network flow, making it '
    'impossible to apply device-specific security policies or to determine whether observed traffic is '
    'consistent with expected device behaviour. Second, behaviour baselining has not been widely '
    'deployed: even where devices are identified, there is no established mechanism to define and '
    'enforce “normal” behaviour on a per-device basis, leaving network operators unable to '
    'detect compromise early enough to prevent damage.')
body_para(doc,
    'This work addresses both problems through a unified framework that (i) classifies device types '
    'from 37 passively observable network flow features using supervised machine learning, and '
    '(ii) detects per-device behavioural anomalies using unsupervised one-class classifiers trained '
    'exclusively on normal device traffic.')

h2(doc, '1.3  Objectives')
first_para(doc, 'The specific objectives of this dissertation are:')
objectives = [
    'To design a set of passively observable, protocol-agnostic network flow features capable of '
    'discriminating between heterogeneous IoT device categories without requiring packet payload '
    'inspection.',
    'To construct a labelled dataset representative of smart home IoT traffic, including injected '
    'anomalous events of multiple attack types.',
    'To train and evaluate multiple classification algorithms for IoT device fingerprinting and select '
    'the most accurate model through empirical comparison on a held-out test set.',
    'To design and implement per-device anomaly detection models based on unsupervised one-class '
    'classification techniques.',
    'To develop a production-grade REST API exposing fingerprinting and anomaly detection as a '
    'service, with automatic schema validation and interactive documentation.',
    'To build a real-time monitoring dashboard for visualising device status and security alerts.',
    'To evaluate the end-to-end system on held-out test data and compare performance against related '
    'published works.',
]
for j, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(style='Body Text')
    set_para_spacing(p, before=0, after=4)
    p.add_run(f'({j})\t').bold = True
    p.add_run(obj)

h2(doc, '1.4  Scope and Delimitations')
body_para(doc,
    'The scope of this dissertation encompasses passive network flow analysis (without packet payload '
    'inspection), eight IoT device types representative of a typical smart home deployment, three '
    'attack categories (data exfiltration, port scanning, and DoS participation), and end-to-end '
    'system implementation from data generation through REST API and dashboard. The framework is '
    'validated on both a synthetic dataset and real traffic samples from the N-BaIoT dataset '
    '(Meidan et al., 2018).')
body_para(doc,
    'The following aspects fall outside the scope of this work: physical layer or firmware analysis; '
    'encrypted payload decryption or protocol reverse engineering; adversarial attacks against the '
    'fingerprinter; live packet capture from a production network; and device types not represented '
    'in the training set. These delimitations are acknowledged and discussed in the context of future '
    'research directions in Chapter 9.')

h2(doc, '1.5  Dissertation Organisation')
body_para(doc,
    'The remainder of this dissertation is structured as follows. Chapter 2 reviews the existing '
    'literature on IoT device fingerprinting, anomaly detection, and machine learning for network '
    'security. Chapter 3 describes the system architecture and design principles. Chapter 4 details '
    'the feature engineering methodology. Chapter 5 covers dataset construction and preprocessing. '
    'Chapter 6 describes the machine learning models employed. Chapter 7 presents the system '
    'implementation, including the REST API and monitoring dashboard. Chapter 8 reports and analyses '
    'experimental results. Chapter 9 concludes the dissertation and identifies directions for future '
    'work. References and appendices follow the main text.')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 2 – LITERATURE REVIEW
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 2: LITERATURE REVIEW')
h2(doc, '2.1  IoT Security: An Overview')
for t in [
    ('The security challenges of IoT ecosystems have been the subject of extensive research since the '
     'proliferation of connected devices accelerated around 2013. The academic literature broadly '
     'categorises these challenges across three layers: the perception layer (physical sensors and '
     'actuators), the network layer (communication protocols and routing), and the application layer '
     '(cloud services and user interfaces).'),
    ('Frustaci et al. (2018) surveyed IoT security challenges across these layers, identifying key '
     'vulnerabilities including weak authentication, insecure update mechanisms, and lack of physical '
     'tamper protection. A comprehensive threat model for smart home environments was proposed by '
     'Apthorpe et al. (2017), who demonstrated through traffic analysis that even encrypted IoT '
     'traffic leaks sensitive behavioural information about occupants. Their findings underscore the '
     'importance of network-level monitoring as a security control.'),
]:
    body_para(doc, t)

h2(doc, '2.2  IoT Device Fingerprinting Techniques')
for t in [
    ('IoT device fingerprinting can be broadly categorised into active and passive approaches. Active '
     'fingerprinting techniques such as Nmap OS detection and banner grabbing are unsuitable for IoT '
     'environments because they require sending probe packets, which may interfere with constrained '
     'devices and reveal the monitoring system to an adversary. Passive fingerprinting, by contrast, '
     'derives device identity solely from observable traffic without generating additional load.'),
    ('Signature-based fingerprinting represents the earliest passive approach. Early work by Bos et '
     'al. (2014) and Radhakrishnan et al. (2014) proposed rule-based signatures derived from DHCP '
     'options, MAC address OUI prefixes, and mDNS records. While effective for known devices, these '
     'methods require manual signature creation and fail to generalise to new device variants.'),
    ('Statistical flow-based fingerprinting emerged as a more scalable alternative. Miettinen et al. '
     '(2017) presented the IoTSentinel system, which used inter-arrival time statistics, byte counts, '
     'and protocol indicators extracted from the first flows following device boot to classify IoT '
     'device types. Their system achieved classification accuracy exceeding 90 % on a corpus of '
     'real IoT devices.'),
    ('Machine learning-based approaches subsequently dominated the literature. Sivanathan et al. '
     '(2018, 2019) conducted seminal work on classifying IoT devices from network traffic using random '
     'forests trained on per-flow features including packet size statistics, DNS query rates, port '
     'numbers, and communication partner diversity. Their system demonstrated that a relatively small '
     'set of passively observable features provides strong discriminative power across diverse device '
     'types.'),
    ('Meidan et al. (2018) proposed an Auto-Encoder-based approach, reconstructing traffic signatures '
     'from auto-encoded latent representations. More recently, Hamza et al. (2019) and Charyyev and '
     'Gunes (2020) demonstrated that graph neural network-based representations of device '
     'communication patterns achieve competitive accuracy, though at significantly higher computational '
     'cost. The present work follows the statistical flow-based approach of Sivanathan et al. (2019), '
     'extending it with an integrated anomaly detection layer.'),
]:
    body_para(doc, t)

h2(doc, '2.3  Anomaly Detection in IoT Networks')
for t in [
    ('Anomaly detection in network traffic has a long history predating IoT, rooted in network '
     'intrusion detection systems (IDS). Traditional signature-based IDS require known attack '
     'signatures and are blind to novel threats. Anomaly-based approaches, by contrast, learn a '
     'model of normal behaviour and flag deviations, enabling detection of previously unseen attacks '
     'at the cost of higher false positive rates.'),
    ('The framing of anomaly detection as a one-class classification problem — training only on '
     'normal data and flagging deviations — is well-established in the literature through methods '
     'such as One-Class SVM (Schölkopf et al., 2001) and deep autoencoders. This formulation is '
     'particularly appropriate for IoT security, where labelled attack data is rarely available for '
     'specific device types and attack patterns evolve rapidly.'),
    ('Liu et al. (2012) proposed the Isolation Forest (IF) algorithm, which explicitly exploits the '
     'property that anomalies are rare and differ from normal instances: anomalous data points are '
     'more easily “isolated” in a random tree partitioning, requiring fewer splits. IF is '
     'computationally efficient, requires no distributional assumptions, and is robust to the curse '
     'of dimensionality, making it well-suited to IoT traffic analysis.'),
    ('In the IoT-specific literature, Meidan et al. (2018) applied deep autoencoders for per-device '
     'anomaly detection, training on benign N-BaIoT traffic and detecting botnet activity through '
     'elevated reconstruction error. Nguyen et al. (2019) demonstrated that behavioural anomaly '
     'detection must be device-type-specific: a single global model trained on all device traffic '
     'fails to capture the dramatically different behavioural envelopes of, for example, a streaming '
     'camera and a constrained motion sensor. This finding directly motivates the per-device '
     'architecture adopted in the present work.'),
]:
    body_para(doc, t)

h2(doc, '2.4  Machine Learning for Network Intrusion Detection')
for t in [
    ('Random Forests (Breiman, 2001) have emerged as one of the most reliable classifiers for network '
     'intrusion detection due to their robustness to irrelevant features, resistance to overfitting '
     'through bagging, and ability to produce calibrated class probability estimates. The '
     'interpretability afforded by Gini importance scores makes them particularly valuable in security '
     'contexts where operator trust requires explainability.'),
    ('Gradient Boosting (Friedman, 2001), particularly XGBoost and LightGBM variants, achieves '
     'competitive or superior accuracy to Random Forests on tabular data by sequentially correcting '
     'residual errors. Support Vector Machines (Cortes and Vapnik, 1995) with RBF kernels provide '
     'strong classification performance through margin maximisation and kernel-induced feature '
     'transformation, though they require careful hyperparameter tuning and do not scale as '
     'gracefully to very large datasets.'),
    ('Ensemble voting classifiers that combine multiple diverse base learners through soft voting '
     '(averaging predicted class probabilities) consistently outperform individual classifiers in '
     'benchmark studies (Fernández-Delgado et al., 2014). This motivates the inclusion of a '
     'soft-voting ensemble as one of the four fingerprinting models evaluated in this work.'),
]:
    body_para(doc, t)

h2(doc, '2.5  Relevant Datasets')
for t in [
    ('The N-BaIoT dataset (Meidan et al., 2018) is the most widely used public dataset for IoT '
     'traffic analysis. It contains benign and attack traffic (Mirai and BASHLITE botnet variants) '
     'captured from nine commercially available IoT devices including a Danmini doorbell, Ecobee '
     'thermostat, and Philips B120N baby monitor. Each flow is described by 115 statistical features '
     'computed over sliding time windows at multiple granularities. N-BaIoT is used in this '
     'dissertation for real-world validation and is integrated through a feature mapping layer that '
     'projects the 115-feature space to the framework’s 37-feature schema.'),
    ('The IoT-23 dataset (Parmisano et al., 2020) captures traffic from 20 malware captures and 3 '
     'benign captures from real IoT devices, with full packet captures available. The UNSW Bot-IoT '
     'dataset (Koroniotis et al., 2019) contains simulated IoT traffic alongside various attack types '
     '(DDoS, DoS, reconnaissance, theft). While these datasets are valuable benchmarks, their device '
     'coverage does not precisely match the eight-device taxonomy adopted in this work. For this '
     'dissertation, a synthetic dataset with statistically grounded per-device profiles was '
     'constructed to ensure class balance and precise control over anomaly injection, with N-BaIoT '
     'providing real-world validation.'),
]:
    body_para(doc, t)

h2(doc, '2.6  Research Gaps and Positioning')
body_para(doc,
    'A review of the literature reveals three key gaps that this dissertation addresses. The first '
    'gap is the lack of unified frameworks: most published systems address either fingerprinting or '
    'anomaly detection, but not both in an integrated pipeline with a shared feature representation '
    'and a common operational interface. The second gap is limited operational deployment: academic '
    'works rarely provide production-grade deployment artefacts — REST APIs, monitoring '
    'dashboards, serialised model artefacts — that would allow practitioners to deploy and '
    'evaluate the system in real environments. The third gap is per-device anomaly specialisation: '
    'many anomaly detection systems use a single global model, ignoring the dramatic behavioural '
    'differences between device types that make a per-device approach essential for high-precision '
    'detection.')
body_para(doc,
    'This work positions itself as a complete, end-to-end, deployable framework that closes all three '
    'gaps, providing both the machine learning core and the operational infrastructure necessary for '
    'real-world deployment.')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 3 – SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 3: SYSTEM ARCHITECTURE AND DESIGN')
h2(doc, '3.1  Design Principles')
body_para(doc,
    'The framework was designed according to five core engineering principles. Modularity ensures that '
    'each functional component — data generation, preprocessing, fingerprinting, anomaly '
    'detection, API, and dashboard — is implemented as an independent Python module with a '
    'well-defined interface, enabling individual components to be replaced or upgraded without '
    'affecting the rest of the system. Passivity requires that no active probing of devices or '
    'network infrastructure occurs; all input is derived from passively collected network flow '
    'records, ensuring the framework is transparent to monitored devices.')
body_para(doc,
    'Scalability is achieved by processing flows as independent records through the preprocessing and '
    'inference pipelines, enabling horizontal scaling. Model serialisation via joblib allows the '
    'trained system to be deployed to new hosts without retraining. Interpretability is supported '
    'through Random Forest Gini importance scores and SHAP (SHapley Additive exPlanations) '
    'attributions computed at inference time, exposing the reasoning behind each prediction to '
    'operators and forensic analysts. Finally, deployability is a design constraint: the entire '
    'system runs on commodity hardware with 8 GB RAM, using exclusively open-source Python '
    'libraries, with no proprietary dependencies.')

h2(doc, '3.2  High-Level Architecture')
body_para(doc,
    'The framework is organised into five logical layers, as illustrated in Figure 3.1. The '
    'Data Ingestion layer (Layer 1) receives 37-feature flow vectors from network monitoring '
    'tools or the synthetic data generator. The Preprocessing layer (Layer 2) applies '
    'RobustScaler normalisation and, during training, SMOTE oversampling. The ML Inference layer '
    '(Layer 3) hosts the device fingerprinter and per-device anomaly detectors, producing '
    'device type labels, confidence scores, and anomaly scores. The REST API layer (Layer 4) '
    'exposes inference capabilities as HTTP endpoints via FastAPI. The User Interface layer '
    '(Layer 5) provides the Plotly Dash monitoring dashboard and the automatically generated '
    'FastAPI Swagger UI.')
add_figure(doc, 'plots/fig3_1_architecture.png',
           'Figure 3.1 — High-Level System Architecture (five logical layers)',
           width=6.0)

h2(doc, '3.3  Module Descriptions')
body_para(doc, 'Table 3.1 provides a summary of the six principal modules and their responsibilities.')
next_table(doc)  # Table 3.1

h2(doc, '3.4  Data Flow')
body_para(doc,
    'The operational data flow for a single incoming network flow proceeds as follows, as illustrated '
    'in Figure 3.2. A 37-value feature vector arrives at the /analyze REST endpoint. The '
    'Preprocessor applies RobustScaler normalisation, producing a scaled 37-dimensional vector. The '
    'DeviceFingerprinter passes the scaled vector through the primary Random Forest model, producing '
    'a device type label and a confidence score; predictions below the 0.75 confidence threshold '
    'are returned as “unknown” rather than a potentially incorrect label. The identified '
    'device type selects the corresponding PerDeviceDetector within the AnomalyDetector. The '
    'PerDeviceDetector computes the Isolation Forest and OC-SVM decision scores, calibrates them to '
    '[0, 1] using power-law scaling, and computes the weighted ensemble score (0.60 ×'
    ' IF + 0.40 × OC-SVM). If the ensemble score is at or above the '
    '0.75 alert threshold, the AlertManager creates an alert with an assigned severity level. The '
    'combined fingerprint and anomaly result is returned in the API response as a JSON object, and '
    'the Dashboard polls the /alerts/recent and /metrics endpoints to update the real-time display.')
add_figure(doc, 'plots/fig3_2_dataflow.png',
           'Figure 3.2 — End-to-End Operational Data Flow',
           width=6.2)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 4 – FEATURE ENGINEERING
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 4: FEATURE ENGINEERING')
h2(doc, '4.1  Feature Selection Rationale')
body_para(doc,
    'The choice of network flow features for IoT device classification must balance discriminative '
    'power with practical observability. Three constraints shaped the feature design. First, '
    'passivity requires that all features be computable from flow-level records (NetFlow v9/IPFIX or '
    'CICFlowMeter output) without access to packet payloads, ensuring the framework operates on '
    'standard network telemetry. Second, protocol agnosticism ensures that features are applicable '
    'regardless of the specific application-layer protocol used, since many IoT protocols are '
    'proprietary or embedded in encrypted TLS streams. Third, temporal stability requires that '
    'features describe steady-state behaviour rather than transient initialisation patterns, enabling '
    'continuous monitoring rather than one-time classification at device boot.')
body_para(doc,
    'A set of 37 features across seven categories was designed through analysis of published IoT '
    'traffic characterisation studies (Sivanathan et al., 2019; Miettinen et al., 2017) and '
    'supplemented with domain knowledge from protocol documentation. The distribution of features '
    'across categories is shown in Figure 4.1.')
add_figure(doc, 'plots/fig4_1_feature_pie.png',
           'Figure 4.1 — Feature Category Distribution (37 Features Across 7 Categories)',
           width=5.5)

h2(doc, '4.2  Feature Categories and Definitions')
body_para(doc, 'Table 4.1 provides the complete definition of all 37 features.')
next_table(doc)  # Table 4.1

h3(doc, 'Category 1: Temporal Features (5 features)')
body_para(doc,
    'Temporal features capture the timing structure of packet inter-arrivals within a flow: '
    'flow_duration, mean_iat, std_iat, min_iat, and max_iat. Different device types exhibit '
    'characteristically different transmission rhythms: streaming cameras emit packets at near-'
    'constant high rates, while event-driven sensors produce bursts separated by long idle periods.')

h3(doc, 'Category 2: Volume Features (4 features)')
body_para(doc,
    'Volume features (packet_count, byte_count, packet_rate, byte_rate) capture the quantity of '
    'network activity. High-bandwidth devices such as cameras and smart TVs are immediately '
    'distinguishable from event-driven sensors by their byte rate alone, which spans more than four '
    'orders of magnitude across device types.')

h3(doc, 'Category 3: Packet Size Features (4 features)')
body_para(doc,
    'Packet size features (mean_pkt_size, std_pkt_size, min_pkt_size, max_pkt_size) capture payload '
    'utilisation patterns. Video streams fill packets to near-MTU values (1,400 bytes), while '
    'command-response devices such as smart plugs use tiny payloads (30–60 bytes). The '
    'standard deviation of packet size further distinguishes devices with uniform payloads from '
    'those with mixed traffic types.')

h3(doc, 'Category 4: Protocol Flag Ratios (6 features)')
body_para(doc,
    'Protocol flag ratios (tcp_ratio, udp_ratio, syn_ratio, fin_ratio, rst_ratio, ack_ratio) '
    'capture the TCP/UDP transport mix and TCP control flag usage patterns, which are highly '
    'characteristic of specific IoT protocols. MQTT-based devices produce predominantly TCP traffic '
    'with a low SYN ratio (persistent connections), while CoAP-based devices use UDP exclusively.')

h3(doc, 'Category 5: Application Layer Features (8 features)')
body_para(doc,
    'Application layer features (is_https, is_mqtt, is_coap, is_mdns, is_ntp, dns_query_count, '
    'well_known_port_ratio, is_encrypted) capture protocol indicators derived from destination port '
    'numbers and packet characteristics. IoT devices are highly protocol-specific: smart bulbs '
    'communicate exclusively via CoAP (port 5683), thermostats via MQTT (port 1883), and smart TVs '
    'via HTTPS (port 443).')

h3(doc, 'Category 6: Traffic Direction Features (3 features)')
body_para(doc,
    'Traffic direction features (upload_bytes, download_bytes, upload_download_ratio) capture the '
    'upload/download asymmetry, which differs fundamentally between device classes. Streaming cameras '
    'exhibit high upload ratios (video upload to cloud), smart TVs exhibit high download ratios '
    '(content streaming), and symmetric devices such as smart speakers exhibit ratios near unity.')

h3(doc, 'Category 7: Destination Diversity Features (7 features)')
body_para(doc,
    'Destination diversity features (unique_dest_ports, unique_dest_ips, port_entropy, ip_entropy, '
    'well_known_ports_count, mean_dest_port, std_dest_port) capture communication breadth. Consumer '
    'devices such as smart TVs contact many cloud endpoints and CDN servers; constrained sensors '
    'communicate with a single hub or gateway. Entropy measures quantify this diversity in a '
    'rotation-invariant manner.')

h2(doc, '4.3  Per-Device Traffic Characteristics')
body_para(doc,
    'The statistical profiles embedded in the synthetic data generator reflect real-world IoT '
    'behaviour characterised in the literature. Table 4.2 summarises the protocol signatures '
    'and key discriminating characteristics for each of the eight device types.')
next_table(doc)  # Table 4.2
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 5 – DATASET
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 5: DATASET CONSTRUCTION AND PREPROCESSING')
h2(doc, '5.1  Synthetic Dataset Generation')
for t in [
    ('A synthetic dataset of 1,600 labelled network flows was constructed using statistically grounded '
     'per-device profiles. The dataset comprises 200 flows per device across 8 device types, of which '
     '190 flows per device represent normal behaviour and 10 flows represent injected anomalies '
     '(5 % anomaly rate). The device distribution is illustrated in Figure 5.1.'),
    ('Synthetic generation was chosen over exclusive reliance on real captured traffic for three '
     'reasons: it allows precise control over class balance, which avoids the severe imbalance '
     'present in real traffic captures; it permits controlled injection of specific attack types at '
     'known rates; and it ensures full reproducibility via a fixed random seed (NumPy seed 42). '
     'Table 5.1 summarises the dataset composition.'),
]:
    body_para(doc, t)
next_table(doc)  # Table 5.1
add_figure(doc, 'plots/fig5_1_dataset_dist.png',
           'Figure 5.1 — Synthetic Dataset Flow Distribution (200 flows per device, 8 devices)',
           width=6.0)
body_para(doc,
    'The generator constructs per-device profiles using clipped normal distributions parameterised '
    'by realistic means and standard deviations derived from the IoT traffic characterisation '
    'literature. For each feature, values are drawn from a clipped normal distribution with hard '
    'bounds enforced via rejection sampling, ensuring that generated feature vectors remain within '
    'physically plausible ranges. A fixed random seed ensures full reproducibility of the dataset '
    'across training runs.')

h2(doc, '5.2  Real-World N-BaIoT Data Integration')
for t in [
    ('To supplement the synthetic dataset and validate generalisation, the framework integrates real '
     'network traffic from the N-BaIoT dataset (Meidan et al., 2018), the most widely used benchmark '
     'for IoT traffic classification. N-BaIoT provides per-device, per-time-window statistical '
     'features for nine commercial IoT devices under both benign and botnet attack conditions.'),
    ('The N-BaIoT features (115 per flow) are mapped to the framework’s 37-feature schema '
     'using a real data loader module (src/data/real_loader.py). Device label mapping assigns '
     'N-BaIoT device identifiers to the nearest corresponding category in the eight-type taxonomy. '
     'The framework supports three dataset construction modes: synthetic (1,600 synthetic flows), '
     'real (N-BaIoT only), and hybrid (N-BaIoT supplemented with synthetic data for device types '
     'not present in N-BaIoT).'),
]:
    body_para(doc, t)

h2(doc, '5.3  Anomaly Injection')
body_para(doc,
    'Three categories of attacks were modelled and injected into the synthetic dataset, as defined '
    'in Table 5.2. Each anomaly is created by multiplying selected feature values by random '
    'factors drawn from specified ranges, producing feature vectors that deviate systematically from '
    'the learned normal profile while remaining within physically plausible bounds. Data exfiltration '
    'is modelled through dramatic increases in upload_bytes and byte_rate. Port scan behaviour is '
    'modelled through extreme increases in unique_dest_ports and port_entropy. DoS participation '
    'is modelled through massive increases in packet_rate and packet_count.')
next_table(doc)  # Table 5.2

h2(doc, '5.4  Preprocessing Pipeline')
body_para(doc,
    'The preprocessing pipeline implements four sequential steps. In Step 1 (Data Cleaning), '
    'infinite values — which can arise from division operations such as the upload/download '
    'ratio with zero download bytes — are replaced with NaN and then imputed with column '
    'medians. All values are clipped to non-negative ranges.')
body_para(doc,
    'Step 2 applies RobustScaler normalisation. RobustScaler is selected over StandardScaler '
    'because it uses the median and interquartile range (IQR) rather than the mean and standard '
    'deviation, rendering it robust to the extreme outliers introduced by anomaly injection. This '
    'ensures that the normalisation of normal flow features is not skewed by the presence of attack '
    'traffic in the dataset.')
body_para(doc,
    'Step 3 applies SMOTE (Synthetic Minority Over-sampling Technique) oversampling. With '
    '190 normal and 10 anomalous flows per device, naive training would produce classifiers biased '
    'toward the majority (normal) class. SMOTE generates synthetic minority samples by linear '
    'interpolation between existing minority instances and their k-nearest neighbours, producing a '
    'balanced training set. The effect of SMOTE on class balance is shown in Figure 5.2.')
add_figure(doc, 'plots/fig5_2_smote.png',
           'Figure 5.2 — SMOTE Oversampling Effect on Class Balance (per device)',
           width=5.8)
body_para(doc,
    'Step 4 performs a stratified train/validation/test split with proportions of 70 %, '
    '15 %, and 15 % respectively. Stratification maintains the class distribution in each '
    'split. The training set is used exclusively for model fitting; the validation set is used for '
    'model selection (selecting the primary fingerprinting model); and the test set is held out '
    'entirely and used only for final performance evaluation, having never been seen during training '
    'or model selection.')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 6 – ML MODELS
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 6: MACHINE LEARNING MODELS')
h2(doc, '6.1  Device Fingerprinting Models')
first_para(doc, 'Four classification algorithms were implemented and evaluated:')

h3(doc, '6.1.1  Random Forest (Primary Model)')
body_para(doc,
    'Random Forest (Breiman, 2001) is an ensemble of decision trees trained on bootstrap samples of '
    'the training data with random feature subsets at each split (the “random subspace” '
    'method). This dual randomisation reduces variance without increasing bias, producing a model '
    'that is robust to overfitting and highly accurate on tabular data. The Random Forest is the '
    'primary production model selected on the basis of validation accuracy. Feature importances '
    'expressed as Gini impurity reduction from the Random Forest are computed and exposed through '
    'the API and dashboard. Hyperparameters are listed in Table 6.1.')
next_table(doc)  # Table 6.1

h3(doc, '6.1.2  Gradient Boosting')
body_para(doc,
    'Gradient Boosting (Friedman, 2001) sequentially adds shallow decision tree weak learners, each '
    'correcting the residual errors of the current ensemble. The model is trained with '
    'n_estimators = 200, learning_rate = 0.1, and max_depth = 4, '
    'providing a strong competitive baseline against the Random Forest.')

h3(doc, '6.1.3  Support Vector Machine (SVM)')
body_para(doc,
    'An RBF-kernel SVM with C = 10, gamma=‘scale’, and '
    'class_weight=‘balanced’ is trained with probability=True to enable soft predictions '
    'through Platt scaling. The balanced class weight setting compensates for class imbalance in the '
    'raw (pre-SMOTE) dataset. SVM provides a complementary decision boundary geometry to the '
    'tree-based methods.')

h3(doc, '6.1.4  Soft Voting Ensemble')
body_para(doc,
    'A VotingClassifier combining Random Forest, Gradient Boosting, and SVM with soft voting '
    '(averaging predicted class probabilities) is trained as the fourth model. Soft voting produces '
    'better-calibrated confidence estimates than hard majority voting and typically outperforms '
    'individual constituent models due to error cancellation across diverse classifiers.')

h3(doc, '6.1.5  Model Selection and Confidence Threshold')
body_para(doc,
    'All four models are trained, and the model achieving the highest validation accuracy is '
    'designated the primary production model for the REST API. A confidence threshold of 0.75 is '
    'applied: predictions for which the maximum class probability falls below this threshold are '
    'returned as “unknown” rather than a potentially incorrect device label. This '
    'conservative threshold prioritises precision over recall in the fingerprinting output, '
    'reflecting the operational preference for flagging ambiguous flows for manual review rather '
    'than making confident but incorrect identifications.')

h2(doc, '6.2  Anomaly Detection Ensemble')
h3(doc, '6.2.1  Architecture: Per-Device Models')
body_para(doc,
    'A critical design decision is the use of per-device anomaly models rather than a single global '
    'model. As demonstrated by Nguyen et al. (2019), a global model trained on all device traffic '
    'fails to capture the dramatically different behavioural envelopes of distinct device types: the '
    'normal traffic of a smart camera (high byte rate, near-MTU packets, sustained TCP upload) is '
    'anomalous relative to the normal traffic of a motion sensor (sporadic CoAP UDP, tiny payloads, '
    'millisecond flows). Training a single model on both conflates their normal envelopes and '
    'reduces detection precision for each individual device type.')
body_para(doc,
    'Accordingly, for each of the eight device types, a separate PerDeviceDetector is trained '
    'exclusively on normal flows for that device type. A global fallback detector, trained on all '
    'normal flows, handles predictions for flows whose identified device type was not seen during '
    'training. The per-device architecture is illustrated in Figure 6.2.')
add_figure(doc, 'plots/fig6_2_anomaly_architecture.png',
           'Figure 6.2 — Per-Device Anomaly Detector Architecture',
           width=6.2)

h3(doc, '6.2.2  Isolation Forest')
body_para(doc,
    'Isolation Forest (Liu et al., 2012) constructs an ensemble of random isolation trees. For each '
    'tree, a random feature and a random split value are selected recursively until each data point '
    'is isolated in its own leaf. Anomalies, being rare and extreme, require fewer splits to isolate '
    'than normal points; the anomaly score is inversely proportional to the average path length '
    'across all trees. The Isolation Forest is trained with n_estimators = 50 and '
    'contamination = 0.05 (matching the 5 % anomaly rate), with random_state ='
    ' 42 for reproducibility.')

h3(doc, '6.2.3  One-Class SVM')
body_para(doc,
    'One-Class SVM (Schölkopf et al., 2001) maps the training data to a high-dimensional '
    'Reproducing Kernel Hilbert Space (RKHS) via the RBF kernel and finds the maximum-margin '
    'hyperplane separating the training data from the origin. The nu parameter (set to 0.05) '
    'controls both the upper bound on the fraction of outliers in training data and the lower bound '
    'on the fraction of support vectors. The kernel bandwidth is set to gamma =‘auto’ '
    '(inverse of n_features), providing a kernel width appropriate to the 37-dimensional feature '
    'space. OC-SVM provides a complementary decision boundary to IF, being effective at capturing '
    'the compact, well-defined behavioural envelope of constrained IoT devices.')

h3(doc, '6.2.4  Score Calibration and Ensemble Weighting')
body_para(doc,
    'Raw decision function outputs from Isolation Forest and OC-SVM are on different scales and '
    'unbounded. A power-law calibration (exponent = 0.6) maps raw scores to [0, 1], '
    'with calibration bounds computed from the 5th and 95th percentile of training set scores. The '
    'upper bound is expanded by 2× to ensure that attack flows produce scores reaching toward '
    '1.0 rather than saturating at the 95th percentile of training data.')
body_para(doc,
    'The ensemble score is the weighted combination: ensemble_score = 0.60 ×'
    ' IF_score + 0.40 × OC-SVM_score. The 60/40 weighting reflects '
    'the empirically observed superiority of Isolation Forest for detecting point anomalies (data '
    'exfiltration, DoS), while OC-SVM contributes more robust boundary detection for structured '
    'behavioural deviations. Alerts are classified into three severity tiers: Medium for scores in '
    '[0.75, 0.85), High for scores in [0.85, 0.95), and Critical for scores of 0.95 '
    'and above. The architecture of the Random Forest fingerprinter is shown in Figure 6.1.')
add_figure(doc, 'plots/fig6_1_rf_architecture.png',
           'Figure 6.1 — Random Forest Decision Ensemble Architecture (100 trees)',
           width=6.0)

h2(doc, '6.3  Model Persistence and Deployment')
body_para(doc,
    'All trained models are serialised using joblib for efficient persistence and fast loading. The '
    'complete set of persisted artefacts includes the fitted RobustScaler, four fingerprinting models '
    '(Random Forest, Gradient Boosting, SVM, and Voting Ensemble), a pointer to the selected primary '
    'model, the 37-value Gini importance array, eight per-device PerDeviceDetectors, and a global '
    'fallback detector. At API startup, all models are loaded into memory via a FastAPI lifespan '
    'context manager, enabling sub-millisecond inference latency for individual flow analysis '
    'requests.')
next_table(doc)  # Table 6.2
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 7 – IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 7: SYSTEM IMPLEMENTATION')
h2(doc, '7.1  Technology Stack')
body_para(doc,
    'Table 7.1 summarises the technology stack employed in this framework. Python 3.9+ '
    'is the primary implementation language. Machine learning components rely on scikit-learn '
    '(classification and anomaly detection), imbalanced-learn (SMOTE), and joblib (model '
    'serialisation). The REST API is implemented with FastAPI and served via Uvicorn. The monitoring '
    'dashboard is built with Plotly Dash. Data manipulation uses NumPy and Pandas; '
    'visualisation uses Matplotlib and Seaborn. Logging is handled by Loguru.')
next_table(doc)  # Table 7.1

h2(doc, '7.2  Project Structure')
body_para(doc,
    'The project follows a modular directory structure under IoT_Device_Fingerprinting_Framework/, '
    'with source code organised under src/ into subpackages for features, data, models, API, and '
    'dashboard. Training is orchestrated by train.py at the project root; the full system is '
    'launched via run.py. Tests reside in the tests/ directory and are executed with pytest.')

h2(doc, '7.3  Training Pipeline')
body_para(doc,
    'The training pipeline (src/models/trainer.py) orchestrates five phases. Phase 1 constructs '
    'the dataset in the selected mode (synthetic, real, or hybrid) and generates the device '
    'distribution visualisation. Phase 2 fits the RobustScaler on the training features, '
    'applies SMOTE to the training set, and produces stratified train/validation/test splits. '
    'Phase 3a trains all four fingerprinting classifiers on the training split and evaluates '
    'each on the validation split to select the primary model. Phase 3b trains per-device '
    'anomaly detectors on normal-only flows for each device type. Phase 4 evaluates all models '
    'on the held-out test set and generates nine evaluation plots (confusion matrices, ROC curves, '
    'feature importances, model comparison bars, and anomaly score distributions). Phase 5 '
    'serialises all trained models and the scaler to the models/ directory.')

h2(doc, '7.4  REST API (FastAPI)')
body_para(doc,
    'The REST API (src/api/main.py) is built with FastAPI, providing automatic OpenAPI documentation '
    'at /docs and asynchronous request handling. All models are loaded at startup via a FastAPI '
    'lifespan context manager. Table 7.2 lists all nine available endpoints.')
next_table(doc)  # Table 7.2
body_para(doc,
    'The POST /explain endpoint accepts a FlowFeatures payload and returns signed SHAP '
    'contributions for the predicted device type. CORS middleware is enabled with permissive '
    'settings to allow the Dashboard and external clients to communicate with the API. Sample '
    'request and response structures for the primary /analyze endpoint are provided in '
    'Appendix B.')

h2(doc, '7.5  Monitoring Dashboard (Plotly Dash)')
body_para(doc,
    'The monitoring dashboard (src/dashboard/app.py) provides a real-time web interface accessible '
    'at the /dashboard/ path. It comprises seven principal panels. The System Status Panel displays '
    'model load state, uptime, total request count, and overall anomaly rate. The Live Alert Feed '
    'provides a scrollable table of recent alerts with timestamp, source IP, device type, anomaly '
    'score, and severity badge. The Anomaly Score Gauge shows the most recent anomaly score against '
    'the 0.75 threshold in real time. The Device Distribution Chart renders a donut chart of '
    'identified device types from recent traffic. The Alert Trend Chart shows a time-series plot '
    'of alert frequency over the past hour. Demo Controls allow injection of simulated attack '
    'traffic for demonstration purposes. The SHAP Explainability Panel renders a live horizontal '
    'bar chart of per-feature SHAP contributions for the most recent prediction, with green bars '
    'indicating features driving the prediction toward the identified device type and red bars '
    'indicating features pulling it away. The dashboard polls the FastAPI backend at a configurable '
    'interval (default: 5 seconds).')

h2(doc, '7.6  Live System Demonstration')
body_para(doc,
    'This section presents a live demonstration of the deployed framework capturing actual system '
    'output during an active session. The demonstration was conducted against the live deployment '
    'at https://mdghufran-iot-fingerprinting.hf.space.')

h3(doc, '7.6.1  FastAPI Swagger UI — Endpoint Overview')
body_para(doc,
    'Figure 7.1 shows the custom dark-themed Swagger UI at the /docs path. The header banner '
    'confirms the live system status: 8 device types, 37 features, 100 Random Forest trees, and '
    'SHAP explainability active. The upper section exposes the GET /health, GET /devices, '
    'POST /fingerprint, and POST /anomaly/score endpoints.')
if screenshot_blobs[0]:
    add_screenshot(doc, screenshot_blobs[0],
        'Figure 7.1 — FastAPI Swagger UI — Live system status banner and upper endpoints')

h3(doc, '7.6.2  FastAPI Swagger UI — Lower Endpoints and Schemas')
body_para(doc,
    'Figure 7.2 shows the lower section of the Swagger UI, revealing the remaining endpoints: '
    'POST /analyze (combined fingerprinting and anomaly scoring), POST /explain (SHAP '
    'explainability), GET /alerts/recent, POST /demo/inject, and the Pydantic request '
    'schema for FlowFeatures with all 37 fields and sensible defaults.')
if screenshot_blobs[1]:
    add_screenshot(doc, screenshot_blobs[1],
        'Figure 7.2 — FastAPI Swagger UI — POST /analyze, /explain and Pydantic schemas')

h3(doc, '7.6.3  API Health Check — Live Response')
body_para(doc,
    'Figure 7.3 presents the JSON response returned by the GET /health endpoint during '
    'the demonstration session. The response confirms status: healthy, models_loaded: true, '
    'shap_ready: true, uptime of 7,551 seconds, and 106 served requests. The 37 ms response '
    'time confirms that the system satisfies the sub-50 ms API latency design requirement.')
if screenshot_blobs[2]:
    add_screenshot(doc, screenshot_blobs[2],
        'Figure 7.3 — GET /health JSON Response — HTTP 200, status: healthy, 37 ms latency')

h3(doc, '7.6.4  Real-Time Dashboard — KPI Strip and Anomaly Timeline')
body_para(doc,
    'Figure 7.4 presents the Plotly Dash dashboard during the active demonstration session. '
    'The KPI strip shows 15 alerts, a 100 % anomaly detection rate on injected flows, '
    '124 total flows processed, and all 8 device types recognised. The live Anomaly Score Timeline '
    'with the 0.75 threshold clearly shows injected attack flows spiking above the threshold.')
if screenshot_blobs[3]:
    add_screenshot(doc, screenshot_blobs[3],
        'Figure 7.4 — Plotly Dash Dashboard — KPI strip and live anomaly timeline')

h3(doc, '7.6.5  Alert Severity Distribution, Score Gauge, and Alerts Table')
body_para(doc,
    'Figure 7.5 shows the lower dashboard panels. The Alerts by Severity bar chart displays '
    'High = 10 and Critical = 6 alerts. The Latest Anomaly Score gauge reads '
    '0.361, placing it in the Normal zone (below the 0.75 threshold), confirming that the system '
    'correctly reports the most recent non-attack flow as benign. The Recent Anomaly Alerts table '
    'lists recent alerts with their device type, score, and severity.')
if screenshot_blobs[4]:
    add_screenshot(doc, screenshot_blobs[4],
        'Figure 7.5 — Dashboard — Alert severity bar chart, anomaly gauge (0.361), alerts table')

h3(doc, '7.6.6  SHAP Live Explainability Panel')
body_para(doc,
    'Figure 7.6 presents the SHAP Live Explainability Panel from the Plotly Dash dashboard. '
    'The panel displays a horizontal bar chart of the top-10 feature contributions for an '
    '“Unknown” prediction at 46 % confidence. Red bars indicate features driving '
    'the prediction away from the identified class; green bars indicate supporting features. The '
    'dominance of tcp_ratio and byte_rate in the SHAP contributions is consistent with the Gini '
    'importance analysis reported in Section 8.3.')
if screenshot_blobs[5]:
    add_screenshot(doc, screenshot_blobs[5],
        'Figure 7.6 — SHAP Live Explainability Panel — Top-10 feature contributions')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 8 – RESULTS
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 8: RESULTS AND EVALUATION')
h2(doc, '8.1  Fingerprinting Performance')
h3(doc, '8.1.1  Test Accuracy and ROC-AUC')
body_para(doc,
    'All four fingerprinting models were evaluated on the held-out test set (15 % of the '
    'SMOTE-oversampled dataset). Results are summarised in Table 8.2 and visualised in '
    'Figure 8.7/8.8. Random Forest, Gradient Boosting, and the Voting Ensemble all achieve '
    'perfect test accuracy (100 %) and a macro ROC-AUC of 1.0000. The SVM achieves '
    '87.08 % accuracy and a macro ROC-AUC of 0.9802, reflecting its higher sensitivity to '
    'feature scale and the separability geometry of the feature space.')
next_table(doc)  # Table 8.2
body_para(doc,
    'Note: Due to the random nature of model selection based on validation accuracy, the primary '
    'model designation may vary between training runs. In practice both Random Forest and the '
    'Voting Ensemble achieve identical test performance; the Random Forest is preferred as primary '
    'because its Gini importances are directly interpretable.')

h3(doc, '8.1.2  Per-Device Classification Report')
body_para(doc,
    'Table 8.1 presents the macro-averaged precision, recall, and F1 score for the Random '
    'Forest model. All device types achieve F1 scores at or above 0.99. Ultra-constrained devices '
    'such as Smart Bulb and Motion Sensor achieve perfect scores due to their extremely distinctive '
    'traffic profiles (CoAP-only, tiny payloads, sub-second flows). Semantically similar devices '
    'such as Smart Speaker and Smart Doorbell, which share HTTPS and MQTT communication patterns, '
    'show the only sub-perfect performance, with a small number of misclassifications between these '
    'two classes.')
next_table(doc)  # Table 8.1

h3(doc, '8.1.3  ROC Curve Analysis')
add_figure(doc, 'plots/fig8_5_roc.png',
           'Figure 8.5 — ROC Curves (One-vs-Rest) — Random Forest',
           width=5.8)
body_para(doc,
    'ROC curves computed using the One-vs-Rest strategy for the Random Forest (Figure 8.5) '
    'reveal near-perfect discrimination for all device types, with per-class AUC values of 1.0000 '
    'across all eight device categories. The curves hug the upper-left corner of the plot, '
    'confirming that the classifier achieves both high true positive rates and very low false '
    'positive rates simultaneously.')

h3(doc, '8.1.4  Confusion Matrix Analysis')
body_para(doc,
    'Confusion matrices for all four models are shown in Figures 8.1 through 8.4. The '
    'Random Forest and Gradient Boosting matrices are diagonal (zero misclassifications). The SVM '
    'matrix reveals misclassifications concentrated between semantically similar device pairs — '
    'specifically Smart Speaker and Smart Doorbell — consistent with their overlapping '
    'application-layer protocol usage. This pattern confirms that the feature engineering '
    'successfully captures the dominant discriminating factors, with remaining ambiguity confined '
    'to device pairs that genuinely exhibit similar network behaviour.')
add_figure(doc, 'plots/fig8_1_cm_random_forest.png',
           'Figure 8.1 — Confusion Matrix — Random Forest (test set)', width=5.5)
add_figure(doc, 'plots/fig8_2_cm_gradient_boosting.png',
           'Figure 8.2 — Confusion Matrix — Gradient Boosting (test set)', width=5.5)
add_figure(doc, 'plots/fig8_3_cm_svm.png',
           'Figure 8.3 — Confusion Matrix — SVM (test set)', width=5.5)
add_figure(doc, 'plots/fig8_4_cm_voting_ensemble.png',
           'Figure 8.4 — Confusion Matrix — Voting Ensemble (test set)', width=5.5)

h2(doc, '8.2  Anomaly Detection Performance')
h3(doc, '8.2.1  Per-Device Anomaly Metrics')
body_para(doc,
    'Table 8.3 reports per-device anomaly detection metrics. All device types achieve ROC-AUC '
    'exceeding 0.96, confirming the effectiveness of the per-device ensemble approach. Smart Bulb '
    'and Motion Sensor achieve the highest scores (AUC > 0.99) due to the extreme '
    'distinctiveness of their normal traffic profiles, which creates very wide separation margins '
    'between normal and anomalous flows.')
next_table(doc)  # Table 8.3
add_figure(doc, 'plots/fig8_9_anomaly_scores.png',
           'Figure 8.9 — Per-Device Anomaly Score Distribution (Normal vs. Attack Traffic)',
           width=6.0)
body_para(doc,
    'Mean anomaly scores for normal traffic (approximately 0.13–0.20 across device types) '
    'are well below the 0.75 alert threshold, while mean scores for attack traffic '
    '(approximately 0.85–0.92) are substantially above it, providing a comfortable detection '
    'margin of approximately 0.65 between the two distributions.')

h3(doc, '8.2.2  Attack Type Analysis')
body_para(doc,
    'The three injected attack types exhibit different levels of detectability. Port scan attacks '
    'are the most easily detected (mean score typically exceeding 0.95) due to extreme values in '
    'unique_dest_ports (500–1,500 vs. normal 1–12) and port_entropy (8–10 bits vs. '
    'normal 0–3 bits), which deviate far outside the learned normal envelope. DoS '
    'participation is highly detectable (mean score 0.88–0.95) through the dramatic packet '
    'rate increase (50–120× normal). Data exfiltration, while well-detectable (mean '
    'score 0.82–0.90) through upload byte and byte rate increases (25–60× normal), '
    'produces slightly lower scores because upload volume alone provides fewer independent signal '
    'dimensions than the multi-feature deviations of port scan behaviour.')

h2(doc, '8.3  Feature Importance Analysis')
h3(doc, '8.3.1  Gini Importance Analysis')
add_figure(doc, 'plots/fig8_6_feature_importance.png',
           'Figure 8.6 — Top-15 Feature Importances (Gini) — Random Forest',
           width=5.8)
body_para(doc,
    'The top-15 most important features by Gini impurity reduction are shown in '
    'Figure 8.6. The distribution of importance across the seven feature categories shows that '
    'volume features contribute approximately 32 % of collective importance, packet size '
    'features approximately 20 %, and protocol flag features approximately 18 %. Byte_rate '
    'is the single most discriminative feature, capturing the fundamental bandwidth difference '
    'between device classes (cameras at 600 KB/s versus motion sensors at 0.05 KB/s). '
    'Protocol boolean flags (is_coap, is_mqtt) exhibit very high per-feature purity for their '
    'associated device classes, reflecting the strict protocol discipline of IoT devices.')

h3(doc, '8.3.2  SHAP Explainability Analysis')
body_para(doc,
    'SHAP (SHapley Additive exPlanations) TreeExplainer is applied to the trained Random Forest on '
    'a stratified sample of 400 flows (50 per device type). The global feature ranking by mean '
    '|SHAP| places tcp_ratio as the strongest fingerprint signal (mean |SHAP| = 0.2386), '
    'reflecting the categorical difference in transport protocol usage between TCP-dominant and '
    'UDP-dominant device classes. The per-device SHAP heatmap reveals that Smart Bulb is the most '
    'uniquely identifiable device, with is_coap dominating its SHAP profile. The SHAP '
    'TreeExplainer is loaded at API startup and exposed via POST /explain, providing '
    'signed per-feature attributions for any incoming flow in real time.')

h2(doc, '8.4  Comparative Analysis')
body_para(doc,
    'Table 8.4 compares the performance of this framework against related published works. '
    'The framework’s fingerprinting performance (100 % accuracy, AUC 1.0000 for '
    'tree-based models) exceeds all published baselines, achieved with a more compact 37-feature '
    'representation compared to the 115-feature N-BaIoT schema. The integrated REST API and '
    'monitoring dashboard constitute a deployment readiness level not matched by any of the '
    'compared academic systems.')
next_table(doc)  # Table 8.4
add_figure(doc, 'plots/fig8_7_8_model_comparison.png',
           'Figure 8.7/8.8 — Model Comparison: Test Accuracy and Macro ROC-AUC',
           width=6.0)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 9 – CONCLUSION
# ════════════════════════════════════════════════════════════════════
h1(doc, 'CHAPTER 9: CONCLUSION AND FUTURE WORK')
h2(doc, '9.1  Summary of Contributions')
body_para(doc,
    'This dissertation has presented, implemented, and evaluated a complete machine learning-based '
    'framework for IoT device fingerprinting and anomaly detection in smart home networks. The '
    'following contributions are made.')
contributions = [
    ('Feature Engineering:',
     'A curated 37-feature set organised into seven semantic categories, carefully designed for '
     'passive, protocol-agnostic IoT traffic analysis. The features span temporal, volume, packet '
     'size, protocol flag, application layer, traffic direction, and destination diversity '
     'characteristics, providing strong discriminative power across eight heterogeneous device types.'),
    ('Dataset:',
     'A statistically grounded synthetic dataset generator with per-device behavioural profiles and '
     'controlled anomaly injection covering three real-world attack types, supplemented by real '
     'traffic integration from the N-BaIoT dataset.'),
    ('Device Fingerprinting:',
     'A multi-algorithm fingerprinting subsystem comparing Random Forest, Gradient Boosting, SVM, '
     'and Voting Ensemble, achieving 100 % test accuracy and macro ROC-AUC of 1.0000 for the '
     'top three models, with Gini importance scores exposed for operator interpretability.'),
    ('SHAP Explainability:',
     'A live SHAP (SHapley Additive exPlanations) explainability layer integrated directly into the '
     'REST API via POST /explain, providing per-prediction, per-feature signed attributions '
     'suitable for forensic investigation and model auditing by security analysts.'),
    ('Anomaly Detection:',
     'A per-device anomaly detection ensemble combining Isolation Forest and One-Class SVM with '
     'power-law score calibration, achieving mean anomaly ROC-AUC exceeding 0.96 across all '
     'device types, with mean attack scores of 0.85–0.92 against a threshold of 0.75.'),
    ('REST API:',
     'A production-grade REST API (FastAPI, port 8000) exposing fingerprinting, anomaly '
     'scoring, combined analysis, SHAP explanation, and alert retrieval endpoints, with automatic '
     'schema validation, interactive Swagger documentation, and sub-50 ms response latency.'),
    ('Dashboard:',
     'A real-time monitoring dashboard (Plotly Dash, port 8050) providing situational '
     'awareness through live alert feeds, device status displays, anomaly score timelines, and '
     'SHAP explainability visualisation.'),
    ('Deployability:',
     'Full reproducibility and deployability on commodity hardware (8 GB RAM, Windows/Linux), '
     'with all source code, data generation scripts, and trained model artefacts publicly available '
     'on Hugging Face Spaces.'),
]
for label, text in contributions:
    p = doc.add_paragraph(style='Body Text')
    set_para_spacing(p, before=2, after=4)
    r = p.add_run(label + '  ')
    r.bold = True
    p.add_run(text)

h2(doc, '9.2  Limitations')
body_para(doc,
    'Several limitations of the current work merit explicit acknowledgement. The synthetic dataset, '
    'while statistically grounded, comprises parameterised distributions rather than live packet '
    'captures; it is possible that real traffic exhibits correlations and transient patterns not '
    'captured by the generator. The closed-world assumption limits the system to classifying devices '
    'into the eight predefined categories: novel device types not represented in training will '
    'receive low-confidence or incorrect fingerprints, requiring periodic model retraining as new '
    'device categories are deployed.')
body_para(doc,
    'The framework is not designed to be robust against adversarial traffic manipulation. A '
    'sophisticated attacker who knows the feature set could craft traffic that mimics a known '
    'benign device’s profile, evading both fingerprinting and anomaly detection. Additionally, '
    'as HTTPS and QUIC adoption increases, even port numbers and IP addresses may become less '
    'reliable discriminators, potentially degrading the information content of several features. '
    'Finally, the synthetic training set of 1,600 flows is sufficient for proof-of-concept '
    'evaluation but smaller than production-scale datasets; real-world deployment would benefit '
    'from a larger and more diverse training corpus.')

h2(doc, '9.3  Future Directions')
body_para(doc,
    'Several directions are identified for future development. Live traffic integration via a '
    'software-defined networking (SDN) controller or lightweight network tap (libpcap/scapy) would '
    'enable the framework to operate on real traffic without requiring pre-processed flow records, '
    'reducing the deployment gap between research and production. Federated learning would enable '
    'training per-home models without sharing raw traffic data across households, addressing privacy '
    'concerns that may limit adoption in residential deployments.')
body_para(doc,
    'Concept drift adaptation through online learning mechanisms would allow the framework to detect '
    'and adapt to gradual changes in device behaviour profiles caused by firmware updates or '
    'behavioural evolution, maintaining detection accuracy without requiring full retraining cycles. '
    'Graph-based device profiling, extending the feature set to include communication graph topology '
    'features such as centrality, clustering coefficient, and device-to-device interaction graphs, '
    'could capture relational anomalies not visible in per-flow analysis.')
body_para(doc,
    'Adversarial robustness through adversarial training (Madry et al., 2018) or certified '
    'defences would harden the fingerprinter against evasion attacks by sophisticated adversaries. '
    'Extended device coverage — expanding the taxonomy from 8 to 20+ device types incorporating '
    'smart locks, medical monitors, baby cameras, and network-attached storage — would broaden '
    'applicability. Deployment on edge hardware such as Raspberry Pi or OpenWRT routers would enable '
    'gateway-level deployment with low latency and no dependency on centralised cloud connectivity, '
    'making the framework viable for isolated smart home networks.')
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════
h1(doc, 'REFERENCES')
refs = [
    ('Apthorpe, N., Reisman, D., & Feamster, N. (2017). A Smart Home is No Castle: Privacy '
     'Vulnerabilities of Encrypted IoT Traffic. Workshop on Data and Algorithmic Transparency.'),
    ('Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.'),
    ('Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic '
     'Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321–357.'),
    ('Charyyev, B., & Gunes, M. H. (2020). IoT Traffic Flow Identification Using Locality-Sensitive '
     'Hashes. IEEE INFOCOM 2020.'),
    ('Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20(3), 273–297.'),
    ('Ericsson. (2023). Ericsson Mobility Report, November 2023. Ericsson AB.'),
    ('Fernández-Delgado, M., Cernadas, E., Barro, S., & Amorim, D. (2014). Do we Need Hundreds '
     'of Classifiers to Solve Real World Classification Problems? Journal of Machine Learning '
     'Research, 15, 3133–3181.'),
    ('Friedman, J. H. (2001). Greedy Function Approximation: A Gradient Boosting Machine. '
     'Annals of Statistics, 29(5), 1189–1232.'),
    ('Frustaci, M., Pace, P., Aloi, G., & Fortino, G. (2018). Evaluating Critical Security Issues '
     'of the IoT World: Present and Future Challenges. IEEE Internet of Things Journal, 5(4), '
     '2483–2495.'),
    ('Hamza, A., Gharakheili, H. H., Benson, T. A., & Sivaraman, V. (2019). Detecting Volumetric '
     'Attacks on IoT Devices via SDN-Based Monitoring of MUD Activity. ACM SOSR 2019.'),
    ('Kolias, C., Kambourakis, G., Stavrou, A., & Voas, J. (2017). DDoS in the IoT: Mirai and '
     'Other Botnets. IEEE Computer, 50(7), 80–84.'),
    ('Koroniotis, N., Moustafa, N., Sitnikova, E., & Turnbull, B. (2019). Towards the development '
     'of realistic botnet dataset in the Internet of Things for network forensic analytics: '
     'Bot-IoT dataset. Future Generation Computer Systems, 100, 779–796.'),
    ('Liu, F. T., Ting, K. M., & Zhou, Z.-H. (2012). Isolation-Based Anomaly Detection. ACM '
     'Transactions on Knowledge Discovery from Data (TKDD), 6(1), 1–39.'),
    ('Lundberg, S. M., & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. '
     'Advances in Neural Information Processing Systems (NeurIPS), 30.'),
    ('Madry, A., Makelov, A., Schmidt, L., Tsipras, D., & Vladu, A. (2018). Towards Deep Learning '
     'Models Resistant to Adversarial Attacks. ICLR 2018.'),
    ('McMahan, H. B., Moore, E., Ramage, D., Hampson, S., & Agüera y Arcas, B. (2017). '
     'Communication-Efficient Learning of Deep Networks from Decentralized Data. AISTATS 2017.'),
    ('Meidan, Y., et al. (2018). N-BaIoT: Network-Based Detection of IoT Botnet Attacks Using '
     'Deep Autoencoders. IEEE Pervasive Computing, 17(3), 12–22.'),
    ('Miettinen, M., et al. (2017). IoT SENTINEL: Automated Device-Type Identification for Security '
     'Enforcement in IoT. IEEE ICDCS 2017.'),
    ('Nguyen, T. D., et al. (2019). Dïot: A Federated Self-learning Anomaly Detection System '
     'for IoT. IEEE ICDCS 2019.'),
    ('Parmisano, A., Garcia, S., & Erquiaga, M. J. (2020). Stratosphere Laboratory — IoT-23: '
     'A labeled dataset with malicious and benign IoT network traffic. Zenodo.'),
    ('Schölkopf, B., Platt, J. C., Shawe-Taylor, J., Smola, A. J., & Williamson, R. C. (2001). '
     'Estimating the Support of a High-Dimensional Distribution. Neural Computation, 13(7), '
     '1443–1471.'),
    ('Sivanathan, A., et al. (2018). Classifying IoT Devices in Smart Environments Using Network '
     'Traffic Characteristics. IEEE Transactions on Mobile Computing, 18(8), 1745–1759.'),
    ('Sivanathan, A., Gharakheili, H. H., & Sivaraman, V. (2019). Managing IoT Cyber-Security '
     'using Programmable Telemetry and Machine Learning. IEEE Transactions on Network and Service '
     'Management, 17(1), 60–74.'),
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.style = doc.styles['Body Text']
    set_para_spacing(p, before=0, after=4)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(-0.5)
    pf.left_indent = Cm(0.5)
page_break(doc)

# ════════════════════════════════════════════════════════════════════
# APPENDICES (copy from source)
# ════════════════════════════════════════════════════════════════════
h1(doc, 'APPENDIX A — FEATURE DEFINITIONS')
h2(doc, 'Complete 37-Feature Vector Specification')
body_para(doc,
    'The following table provides the complete specification of all 37 network flow features used '
    'by the framework. Features are listed in the order in which they appear in the feature vector '
    'passed to the preprocessing pipeline and model inference endpoints.')
# copy Appendix A table if available (it appears after the 13 main tables)
# Tables 0-12 are the 13 main tables; remaining are appendix tables
for _ in range(3):  # three appendix tables
    if TABLE_IDX < len(src_tables):
        next_table(doc)

page_break(doc)
h1(doc, 'APPENDIX B — API ENDPOINT REFERENCE')
h2(doc, 'Base URL')
body_para(doc, 'https://mdghufran-iot-fingerprinting.hf.space')

endpoints = [
    ('GET /health',
     'Returns service health status including model load state, SHAP readiness, uptime, and '
     'total request count.'),
    ('GET /devices',
     'Returns the list of eight supported device types with descriptions.'),
    ('POST /fingerprint',
     'Identifies device type from a 37-feature flow vector. Returns device_type, confidence, '
     'is_known, and model_used.'),
    ('POST /anomaly/score',
     'Scores a flow against the per-device anomaly baseline. Returns anomaly_score, is_anomalous, '
     'severity, and threshold.'),
    ('POST /analyze',
     'Combined fingerprinting and anomaly detection in a single request. Returns the full '
     'fingerprint and anomaly result objects.'),
    ('GET /alerts/recent?n=50',
     'Returns the last N alerts. Default: 50.'),
    ('GET /metrics',
     'Returns aggregate metrics: uptime, request count, total alerts, and anomaly rate.'),
    ('POST /explain',
     'Returns SHAP-based signed feature attributions for the predicted device type. Accepts '
     'a FlowFeatures payload (37 features).'),
    ('POST /demo/inject',
     'Injects a simulated attack flow for demonstration purposes. Used by the dashboard '
     'Demo Controls panel.'),
]
for ep_name, ep_desc in endpoints:
    h3(doc, ep_name)
    body_para(doc, ep_desc)

page_break(doc)
h1(doc, 'APPENDIX C — SYSTEM REQUIREMENTS AND SETUP')
h2(doc, 'Hardware Requirements')
body_para(doc,
    'Minimum: 8 GB RAM, dual-core CPU, 2 GB free disk space. Recommended: 16 GB RAM '
    'for comfortable operation with SHAP explainability enabled. The system runs on Windows 10+ '
    'and Linux (Ubuntu 20.04+).')

h2(doc, 'Software Setup')
h3(doc, 'Step 1: Install Python')
body_para(doc,
    'Python 3.9 or higher is required. Verify installation with: python --version')

h3(doc, 'Step 2: Install Dependencies')
body_para(doc, 'Execute: pip install -r requirements.txt')
body_para(doc,
    'Required packages include: numpy, pandas, scikit-learn, imbalanced-learn, joblib, fastapi, '
    'uvicorn, pydantic, dash, plotly, loguru, matplotlib, seaborn, and pytest.')

h3(doc, 'Step 3: Train the Models')
body_para(doc, 'Execute: python train.py')
body_para(doc,
    'This command generates the synthetic 1,600-flow dataset, preprocesses it with RobustScaler '
    'and SMOTE, trains all four fingerprinting models and the eight per-device anomaly detectors, '
    'generates evaluation plots, and serialises all artefacts to the models/ directory. Expected '
    'training time: 2–5 minutes on a modern laptop.')

h3(doc, 'Step 4: Launch the Framework')
body_para(doc, 'Execute: python run.py')
body_para(doc,
    'This starts both the FastAPI server (port 8000) and the Plotly Dash dashboard '
    '(port 8050). The API documentation is accessible at /docs and the dashboard at /dashboard/.')

h3(doc, 'Step 5: Run Tests')
body_para(doc, 'Execute: pytest tests/ -v')

h2(doc, 'Running on Real N-BaIoT Data')
body_para(doc,
    'Place the N-BaIoT CSV files in data/nbaiot/ following the project directory structure. '
    'Then run: python train.py --mode real (for N-BaIoT only) or python train.py --mode hybrid '
    '(to supplement missing device types with synthetic data).')

doc.add_paragraph()
p = doc.add_paragraph('End of Dissertation')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=36, after=12)
p.runs[0].bold = True; p.runs[0].font.size = Pt(12)

p2 = doc.add_paragraph(
    'Md Ghufran Alam  |  Roll No. NDU202400038\n'
    'M.Tech Cyber Forensics, Session 2024–2026\n'
    'NIELIT Srinagar, Jammu & Kashmir  |  May 2026')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p2, before=0, after=0)

# ════════════════════════════════════════════════════════════════════
doc.save(DEST)
print(f'Saved: {DEST}')
print(f'Tables copied from source: {TABLE_IDX} of {len(src_tables)}')
