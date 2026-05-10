"""
Generates a Word document (.docx) containing the full panel presentation script
for IoT Device Fingerprinting & Anomaly Detection Framework.
Output: Presentation_Script.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)


# ── Style helpers ────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False,
             color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_paragraph(para, hex_color="1A1E28"):
    """Apply cell-like background shading to a paragraph via XML."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_page_break():
    doc.add_page_break()


def title_page():
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRESENTATION SCRIPT")
    set_font(r, size=22, bold=True, color=(0, 180, 200))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "Design of Framework for IoT Device Fingerprinting and\n"
        "Anomaly Detection for Smart Home Using Machine Learning"
    )
    set_font(r2, size=15, bold=True)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Md Ghufran Alam  |  NDU202400038")
    set_font(r3, size=13, bold=True)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("M.Tech Cyber Forensics  |  NIELIT Srinagar  |  Batch 2024–2026")
    set_font(r4, size=12, italic=True, color=(100, 100, 100))

    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(
        "This document contains the full word-by-word speaking script\n"
        "for your panel presentation. Read it naturally — speak slowly and clearly."
    )
    set_font(r5, size=11, italic=True, color=(120, 120, 120))

    doc.add_paragraph()

    # Divider
    p6 = doc.add_paragraph("─" * 70)
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # How to use box
    p7 = doc.add_paragraph()
    r7 = p7.add_run("HOW TO USE THIS SCRIPT")
    set_font(r7, size=12, bold=True, color=(0, 180, 200))

    tips = [
        "Each slide has: WHAT TO SAY (your spoken words) + KEY POINTS TO REMEMBER.",
        "Words inside [ square brackets ] are actions — do NOT say them aloud.",
        "Speak at a calm pace. One slide ≈ 1.5 to 2 minutes.",
        "Total presentation target: 20–25 minutes + 5–10 minutes Q&A.",
        "If a panelist interrupts, say: 'Great question — I will cover that shortly' or answer directly.",
        "Maintain eye contact. Glance at slides, but speak TO the panel, not AT the slide.",
    ]
    for t in tips:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t)
        set_font(r, size=11)

    add_page_break()


def slide_section(number, title, duration="~1.5 min"):
    """Adds a slide heading block."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"  SLIDE {number}  —  {title.upper()}  ({duration})  ")
    set_font(r, size=14, bold=True, color=(255, 255, 255))
    shade_paragraph(p, "0D4F6E")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def what_to_say(text):
    """Main spoken script block."""
    p = doc.add_paragraph()
    label = p.add_run("WHAT TO SAY:\n")
    set_font(label, size=10, bold=True, color=(0, 160, 180))

    r = p.add_run(text.strip())
    set_font(r, size=12)
    doc.add_paragraph()


def key_points(points):
    """Key technical points to remember."""
    p = doc.add_paragraph()
    r = p.add_run("KEY POINTS TO REMEMBER:")
    set_font(r, size=10, bold=True, color=(160, 80, 220))
    shade_paragraph(p, "F5F0FF")

    for pt in points:
        p2 = doc.add_paragraph(style="List Bullet")
        r2 = p2.add_run(pt)
        set_font(r2, size=11)

    doc.add_paragraph()


def likely_questions(qas):
    """Panel Q&A preparation."""
    p = doc.add_paragraph()
    r = p.add_run("LIKELY PANEL QUESTIONS & YOUR ANSWERS:")
    set_font(r, size=10, bold=True, color=(180, 80, 0))
    shade_paragraph(p, "FFF5EE")

    for q, a in qas:
        pq = doc.add_paragraph()
        rq = pq.add_run(f"Q: {q}")
        set_font(rq, size=11, bold=True, color=(180, 80, 0))

        pa = doc.add_paragraph()
        ra = pa.add_run(f"A: {a}")
        set_font(ra, size=11)

    doc.add_paragraph()


def transition(text):
    p = doc.add_paragraph()
    r = p.add_run(f"[ TRANSITION → {text} ]")
    set_font(r, size=10, italic=True, color=(140, 140, 140))
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════

title_page()

# ── SLIDE 1 ─────────────────────────────────────────────────────────────────
slide_section(1, "Title Slide", "~1 min")

what_to_say("""
Good morning, respected panel members.

My name is Md Ghufran Alam. My roll number is NDU202400038.
I am a student of M.Tech Cyber Forensics at NIELIT Srinagar, currently in my final semester.

Today, I am presenting my thesis project titled:

"Design of a Framework for IoT Device Fingerprinting and Anomaly Detection
for Smart Home using Machine Learning."

In simple words — I have built a system that can automatically identify what type
of smart home device is on a network, just by looking at its traffic patterns.
And if any device starts behaving in a suspicious way — like sending too much data
or scanning other devices — my system detects it and raises an alert.

I will walk you through the problem, my approach, the models I built, and the results.
Thank you for giving me this opportunity.
""")

key_points([
    "Project: IoT fingerprinting + anomaly detection",
    "Your name: Md Ghufran Alam | Roll: NDU202400038",
    "Programme: M.Tech Cyber Forensics, NIELIT Srinagar",
    "Core idea: Identify device type from traffic, detect anomalies",
])

transition("Problem Statement")

# ── SLIDE 2 ─────────────────────────────────────────────────────────────────
slide_section(2, "Problem Statement", "~2 min")

what_to_say("""
Let me first explain WHY this problem is important.

By the year 2030, there will be more than 30 BILLION IoT devices in the world.
These devices include smart cameras, smart bulbs, smart TVs, smart doorbells —
basically everything in a modern home is now connected to the internet.

But here is the serious problem: around 70% of these IoT devices have NO proper
security built into them. They don't have unique identities. They don't have
strong authentication. Once they are connected to a home network, anyone can
pretend to be them, or they can be hijacked by attackers.

The attack surface — meaning the number of ways an attacker can enter a network —
has grown by 10 times in the last few years, just because of IoT devices.

Now, what kind of attacks happen?
— An attacker can secretly make a smart camera upload all your video data to an
  outside server. This is called Data Exfiltration.
— A compromised smart bulb can scan all your home devices to find which ones
  are open. This is called a Port Scan.
— Multiple smart devices can be made to flood a target with traffic. This is called
  a DoS attack — Denial of Service.

The biggest problem is: existing security systems were NOT designed for IoT.
Traditional firewalls and antivirus systems do not understand IoT device behaviour.
There is no per-device baseline — so nobody knows what is normal for a smart bulb
versus a smart camera.

My project solves exactly this problem.
""")

key_points([
    "30 billion IoT devices by 2030 — most have zero security",
    "No built-in identity for IoT = rogue devices go undetected",
    "3 attack types: Data Exfiltration, Port Scan, DoS Participation",
    "Traditional IDS systems not designed for constrained IoT devices",
    "Per-device baseline is missing — can't tell normal from abnormal",
])

likely_questions([
    ("Why is IoT security different from regular network security?",
     "IoT devices are resource-constrained — they cannot run antivirus or firewalls. "
     "They use unusual protocols like MQTT and CoAP. Traditional systems don't "
     "understand their traffic patterns, so we need specialized ML-based approaches."),
    ("What do you mean by passive fingerprinting?",
     "Passive means we only OBSERVE the traffic — we don't interfere with it, "
     "we don't send any packets to the device. We just analyze what the device "
     "sends naturally. This is non-intrusive and works in real-time."),
])

transition("Objectives")

# ── SLIDE 3 ─────────────────────────────────────────────────────────────────
slide_section(3, "Objectives", "~1.5 min")

what_to_say("""
Based on this problem, I defined 6 clear objectives for my project.

Objective 1: Design a passive, non-intrusive machine learning framework
for IoT device fingerprinting. This means no deep packet inspection —
I don't look inside the packets, I only look at traffic-level patterns.

Objective 2: Extract 37 meaningful features from raw network flows.
These features cover timing, volume, packet size, protocols, and more.

Objective 3: Classify 8 different smart home device types with more than
95% accuracy using machine learning.

Objective 4: Detect three types of anomalies — data exfiltration, port scanning,
and DoS participation — using a per-device ensemble model.

Objective 5: Deploy a production-ready REST API using FastAPI so that this
system can be used in real applications, not just as a research prototype.

And Objective 6: Build per-device behavioural baselines, so that each device
has its own normal profile. This greatly reduces false positives.

All 6 objectives have been successfully achieved, as I will show in the results.
""")

key_points([
    "6 objectives — all achieved",
    "Passive = no DPI, no packet content inspection",
    "37 features extracted from traffic metadata only",
    "8 device types classified with 100% accuracy",
    "Per-device baselines eliminate false positives",
])

transition("System Architecture")

# ── SLIDE 4 ─────────────────────────────────────────────────────────────────
slide_section(4, "System Architecture", "~2.5 min")

what_to_say("""
Now let me explain the overall architecture of my framework.

Think of it as a pipeline — data flows from one step to the next,
just like an assembly line.

STEP 1 — Data Generation:
Since I don't have access to a real smart home lab with 8 different IoT devices,
I created a synthetic dataset. This dataset has 1,600 network flows — 200 flows
per device type — each described by 37 statistical features. Each device type
has a statistically distinct traffic profile, so the data is realistic.

STEP 2 — Feature Extraction:
From each network flow, I extract 37 features. These include things like
packet rate, byte count, protocol ratios like TCP vs UDP, application-layer
flags like MQTT and CoAP, and destination entropy. I'll explain these in detail
on the next slide.

STEP 3 — Preprocessing:
I apply RobustScaler to normalize all features. I specifically chose RobustScaler
instead of StandardScaler because it is resistant to outliers — which is important
for IoT traffic that can have extreme values during attacks.
I also apply SMOTE — Synthetic Minority Oversampling Technique — to balance
the class distribution in the data.

STEP 4 — Device Fingerprinting:
I train a Random Forest classifier with 100 trees to identify which device type
the traffic belongs to. I also trained Gradient Boosting and SVM for comparison.

STEP 5 — Anomaly Detection:
For each device type, I train a separate ensemble of Isolation Forest
and One-Class SVM. This learns what "normal" looks like for that specific device.

STEP 6 — Alert Manager, API, and Dashboard:
Anomalies are routed to an alert manager that assigns severity levels.
Everything is exposed through a FastAPI REST API, and a Plotly Dash
dashboard provides live visual monitoring.

So the full flow is: Raw Traffic → Features → Models → Alerts → API → Dashboard.
""")

key_points([
    "6-stage pipeline: Data → Features → Preprocess → Fingerprint → Anomaly → Output",
    "API on Port 8000 (FastAPI), Dashboard on Port 8050 (Plotly Dash)",
    "RobustScaler normalizes features (outlier-resistant); SMOTE balances class distribution",
    "Per-device anomaly models — not one global model",
])

likely_questions([
    ("Why did you use synthetic data instead of real traffic?",
     "Real IoT traffic capture requires physical devices, a controlled lab environment, "
     "and ethical approvals. Synthetic data lets me create statistically valid, "
     "labelled data for all 8 device types. The feature distributions are based on "
     "real-world IoT traffic studies. In future work, I plan to validate with real captures."),
    ("What is SMOTE and why did you need it?",
     "SMOTE is Synthetic Minority Oversampling Technique. Since all 8 device classes "
     "have equal samples, the classes are already balanced. But I applied SMOTE as a "
     "precautionary step after the train-test split to ensure the training set is "
     "not biased. It creates synthetic samples by interpolating between existing ones."),
])

transition("Dataset Description")

# ── SLIDE 5 ─────────────────────────────────────────────────────────────────
slide_section(5, "Dataset Description", "~2 min")

what_to_say("""
Let me describe the dataset I used.

I generated 1,600 network flow records in total.
Each of the 8 IoT device types has exactly 200 flows.

Out of these 200 flows per device, 190 are normal flows, and 10 are anomalous —
that is a 5% anomaly fraction, which is realistic for real-world network traffic.

The 8 device types I cover are:
Smart Camera, Smart Thermostat, Smart TV, Smart Bulb, Smart Plug,
Smart Speaker, Smart Doorbell, and Motion Sensor.

Each device has a very distinct traffic signature.
For example — a Smart Camera generates about 40 megabytes per flow because
it streams continuous video. But a Motion Sensor generates only about 80 bytes
per flow because it just sends a tiny event trigger when motion is detected.
A Smart TV generates even more than the camera — around 700 megabytes per flow
because it streams 4K video from cloud services.

These extreme differences make the fingerprinting problem very tractable,
and my model achieves 100% accuracy because the device signatures are so distinct.

For splitting the data, I used a 70-15-15 split:
70% for training, 15% for validation, and 15% for testing.
RobustScaler was fit on the full dataset and then used to transform all splits.
RobustScaler uses median and IQR instead of mean and std, so extreme anomaly
values do not distort the scaling.
""")

key_points([
    "1,600 flows total — 200 per device, 5% anomaly (10 flows/device)",
    "8 device types with extremely distinct traffic signatures",
    "Smart Camera = ~40MB/flow vs Motion Sensor = ~80 bytes/flow",
    "70 / 15 / 15 split — RobustScaler (median+IQR based, outlier-resistant)",
    "3 anomaly types: Data Exfiltration, Port Scan, DoS Participation",
])

transition("Feature Engineering")

# ── SLIDE 6 ─────────────────────────────────────────────────────────────────
slide_section(6, "Feature Engineering", "~2 min")

what_to_say("""
Feature engineering is the heart of this project.

I extract 37 features from each network flow, organized into 7 categories.

Category 1 — TEMPORAL features — 5 features:
These capture the timing of packets. For example, flow_duration tells us
how long a flow lasted. Mean_iat is the average time gap between consecutive packets.
A camera has very short inter-arrival times because it sends packets continuously.
A thermostat has very long inter-arrival times because it only sends data every few minutes.

Category 2 — VOLUME features — 4 features:
These capture how much data is being sent. Packet count, byte count,
packet rate, and byte rate. A Smart TV has a massive byte rate.
A Smart Bulb has almost zero.

Category 3 — PACKET SIZE features — 4 features:
Mean, standard deviation, minimum, and maximum packet size.
Cameras and TVs send large packets — close to 1,500 bytes — the maximum.
Sensors send tiny packets — just 20 to 50 bytes.

Category 4 — PROTOCOL FLAGS — 6 features:
TCP ratio, UDP ratio, SYN ratio, FIN ratio, RST ratio, ACK ratio.
MQTT devices like thermostats use mostly UDP. HTTPS cameras use mostly TCP.

Category 5 — APPLICATION LAYER — 8 features:
These are binary flags: is_https, is_mqtt, is_coap, is_mdns, is_ntp,
plus dns_query_count, well_known_port_ratio, and is_encrypted.
This category alone can often distinguish a device type by itself.

Category 6 — TRAFFIC DIRECTION — 3 features:
Upload bytes, download bytes, and their ratio.
A camera uploads more than it downloads. A TV downloads much more than it uploads.

Category 7 — DESTINATION features — 7 features:
Unique destination ports, unique destination IPs, port entropy, IP entropy,
and related features. A port scan will show extreme values here — 500+ unique ports
with very high entropy.

The key insight is: NO device has the same combination across all 37 features.
That is what makes fingerprinting possible.
""")

key_points([
    "37 features across 7 categories — all from traffic metadata, no payload",
    "Temporal (5): flow_duration, mean_iat, std_iat, min_iat, max_iat",
    "Volume (4): packet_count, byte_count, packet_rate, byte_rate",
    "Protocol Flags (6): tcp/udp/syn/fin/rst/ack ratios",
    "Application Layer (8): is_https, is_mqtt, is_coap, is_encrypted, etc.",
    "Destination (7): unique_dest_ports, port_entropy, ip_entropy — key for anomaly",
    "No Deep Packet Inspection — only header-level information",
])

likely_questions([
    ("Why 37 features specifically? How did you choose them?",
     "I reviewed IoT traffic analysis literature and identified features that are "
     "computationally cheap to extract but highly discriminative across device types. "
     "The 7 categories cover all aspects of flow behaviour — timing, volume, size, "
     "protocol, application, direction, and destination. Together they create a unique "
     "fingerprint for each device type."),
    ("Can these features be extracted in real-time?",
     "Yes. All 37 features are flow-level statistics that can be computed within "
     "a few seconds of a flow completing. Tools like CICFlowMeter or custom scripts "
     "on a network tap or gateway router can extract these in real time."),
])

transition("Device Fingerprinting Module")

# ── SLIDE 7 ─────────────────────────────────────────────────────────────────
slide_section(7, "Device Fingerprinting Module", "~2.5 min")

what_to_say("""
Now let me explain the Device Fingerprinting Module.

The task here is: given the 37 features of a network flow, classify it into
one of the 8 device types.

I chose RANDOM FOREST as my primary model. Here is why:

Random Forest is an ensemble of decision trees. I used 100 trees with
a maximum depth of 15 levels. It is robust to outliers, handles class
imbalance well with the class_weight='balanced' setting, and works
very well with tabular data like network features. It also gives us
feature importances, which helps us understand which features matter most.

I also set a CONFIDENCE THRESHOLD of 0.75.
When a model makes a prediction, it gives a confidence score — a probability.
If the confidence is below 75%, the system reports the device as "unknown"
instead of guessing. This avoids wrong classifications when a device is new
or behaves unusually.

For comparison, I also trained:
— Gradient Boosting: 100 trees, learning rate 0.1
— SVM with RBF kernel: C=10, probability estimates enabled
— A Voting Ensemble that combines all three using soft voting

The best model is automatically selected based on validation accuracy.
In my experiments, Random Forest and the Voting Ensemble both achieved
100% accuracy on the test set. Random Forest was selected as primary
because it is faster and more interpretable.

The test accuracy of 100% is because the 8 device types have very
distinct statistical signatures. The features separate them very cleanly.
""")

key_points([
    "Primary: Random Forest — 100 trees, max_depth=15, class_weight=balanced",
    "Confidence threshold 0.75 — below this → device reported as 'unknown'",
    "4 models trained: RF, Gradient Boosting, SVM (RBF), Voting Ensemble",
    "Best model auto-selected by validation accuracy",
    "Result: 100% test accuracy, ROC-AUC = 1.000",
])

likely_questions([
    ("Why not use a deep learning model like a neural network?",
     "For tabular data with 37 well-engineered features, Random Forest consistently "
     "outperforms neural networks. Neural networks need much more data and are harder "
     "to interpret. Random Forest is also more suitable for deployment on edge devices "
     "with limited resources. Deep learning is listed in my future work."),
    ("What does ROC-AUC of 1.000 mean?",
     "ROC-AUC measures how well a model separates classes. A score of 1.0 means "
     "perfect separation — the model can perfectly distinguish every device type "
     "from every other. This is because the device traffic profiles are very distinct."),
    ("Is 100% accuracy realistic? Isn't that overfitting?",
     "The 100% is on the TEST set, which the model never saw during training. "
     "The high accuracy reflects the fact that the device traffic signatures are "
     "genuinely very different — a camera sending 40MB vs a sensor sending 80 bytes "
     "is easy to separate. On real noisy traffic, we would expect 95%+ accuracy."),
])

transition("SHAP Feature Importance — Model Explainability")

# ── SLIDE 8 ─────────────────────────────────────────────────────────────────
slide_section(8, "SHAP Feature Importance", "~1.5 min")

what_to_say("""
Before I move to anomaly detection, I want to address a very important concern
that experts always raise about machine learning models — the BLACK BOX problem.

When a model says "this device is a smart camera", can we trust it?
Does it actually use meaningful features, or has it learned some spurious pattern?

I solved this using SHAP — SHapley Additive exPlanations.
SHAP is a game-theory based technique that quantifies exactly how much
each feature contributed to each individual prediction.

This is NOT just a global average. SHAP gives a separate explanation
for every single prediction the model makes.

Looking at the bar chart on screen, these are the top 15 features ranked by
their average absolute SHAP value across all 400 test samples:

Number 1: tcp_ratio — whether the device uses TCP. This makes sense.
  Cameras, TVs, and doorbells are almost entirely TCP.
  Bulbs and sensors are almost entirely UDP.

Number 2: udp_ratio — the flip side of tcp_ratio.

Number 3: ack_ratio — the fraction of ACK packets.
  TCP-heavy devices have many more ACK packets.
  This distinguishes cameras from sensors perfectly.

Number 4: mean_dest_port — the average destination port.
  Port 443 means HTTPS — cameras and TVs.
  Port 1883 means MQTT — thermostats and plugs.
  Port 5683 means CoAP — smart bulbs.

Number 5: is_mqtt — application layer flag.
  When this is 1, the device is almost certainly a thermostat or plug.

The SHAP values prove that my model is using exactly the RIGHT features —
the features that actually distinguish device types by their protocol behaviour.
This is a very strong argument that the model will generalize to real traffic.

I also built a LIVE SHAP demo — the POST /explain endpoint in my API
returns these SHAP values in real time for any incoming flow.
And the live dashboard shows a green/red bar chart updating every 5 seconds.
""")

key_points([
    "SHAP = game-theory based explainability — not global averages, per-prediction",
    "#1 tcp_ratio, #2 udp_ratio — protocol choice is the strongest fingerprint",
    "#3 ack_ratio — distinguishes TCP-heavy cameras from UDP-heavy sensors",
    "#4 mean_dest_port — port 443=camera/TV, port 1883=thermostat, port 5683=bulb",
    "#5 is_mqtt — alone identifies thermostat/plug with high confidence",
    "POST /explain API endpoint returns live SHAP for any network flow",
])

likely_questions([
    ("What is SHAP and why did you choose it over LIME or other methods?",
     "SHAP is based on Shapley values from cooperative game theory. It is the only "
     "method that satisfies three mathematical properties: efficiency (values sum to "
     "the prediction), symmetry (same features get same credit), and null player "
     "(features that don't contribute get zero). LIME only gives local approximations "
     "and can be inconsistent. SHAP is also natively supported for tree models like "
     "Random Forest through TreeExplainer, making it very fast."),
    ("Why is tcp_ratio the most important feature? Isn't that too simple?",
     "Simple features can be the most discriminative. A smart bulb using CoAP has "
     "nearly 0% TCP, while a smart camera streaming video has 92% TCP. This single "
     "feature already splits 8 device types into two clear groups. The model is "
     "right to use it — it is a genuine, stable characteristic of device behaviour, "
     "not a spurious correlation."),
    ("Can SHAP be fooled by an adversary who mimics protocol patterns?",
     "Yes, protocol mimicry is a known attack. If an attacker makes malicious traffic "
     "look like normal camera traffic at the protocol level, it would fool the model. "
     "This is why I combine fingerprinting with anomaly detection — even if the "
     "protocol pattern is mimicked, the byte count or port entropy might still be "
     "anomalous enough to trigger an alert. Defense in depth."),
])

transition("Per-Device SHAP Heatmap")

# ── SLIDE 9 ─────────────────────────────────────────────────────────────────
slide_section(9, "Per-Device SHAP Heatmap", "~1 min")

what_to_say("""
This heatmap shows a different view of SHAP — per device type.

Each row is one of the 8 device types.
Each column is one of the top 12 most important features.
The brightness of each cell shows how much that feature influences
the prediction for that specific device type.
A bright cell means: this feature is CRITICAL for identifying this device.

Let me walk through the key patterns:

Smart Camera: tcp_ratio and is_https are very bright — the camera uses HTTPS
for video streaming. ack_ratio is also bright because TCP streams have many ACKs.

Smart Bulb: is_coap completely dominates its row — almost every other cell is dark.
CoAP on port 5683 is the sole identifier for a smart bulb.
This means: if you see CoAP, it is almost certainly a bulb. Certainty near 100%.

Smart Thermostat: is_mqtt and mean_dest_port (1883) are the dominant features.
The thermostat periodically reports temperature data via MQTT broker.

Motion Sensor: Its brightest feature is byte_count — specifically, its VERY LOW
byte count. A motion sensor sends tiny event packets when motion is detected.
The model learned that near-zero byte count is the motion sensor signature.

This heatmap is extremely useful for explaining my model to a non-technical
audience — you can literally point to a cell and say "this is why the model
thinks that device is a smart bulb."
""")

key_points([
    "Rows = 8 device types, Columns = top 12 features, Brightness = SHAP influence",
    "Smart Bulb: is_coap alone dominates — CoAP port 5683 is a unique separator",
    "Smart Camera: tcp_ratio + is_https + ack_ratio form the fingerprint",
    "Motion Sensor: near-zero byte_count is its defining characteristic",
    "Thermostat: is_mqtt + mean_dest_port 1883 completely identify it",
])

likely_questions([
    ("What if two devices have similar SHAP patterns?",
     "Smart plug and thermostat are the hardest to distinguish — both use MQTT "
     "on port 1883. The differentiating features are byte_count and packet_rate: "
     "the thermostat sends periodic small temperature readings, while the smart plug "
     "has slightly different power-reporting patterns. The model still achieves 100% "
     "accuracy on both, but with slightly lower confidence margins between these two."),
])

transition("Anomaly Detection Module")

# ── SLIDE 10 ─────────────────────────────────────────────────────────────────
slide_section(10, "Anomaly Detection Module", "~2.5 min")

what_to_say("""
Now let me explain the Anomaly Detection Module — this is the more complex
and novel part of my project.

The challenge with anomaly detection is that we cannot define what an anomaly
looks like in advance. Instead, we learn what NORMAL looks like — and then
anything that deviates significantly from normal is flagged as anomalous.

My key design decision was: USE PER-DEVICE MODELS, not one global model.

Think about it — a smart camera normally sends 40 megabytes per flow.
If a smart bulb suddenly sends 10 megabytes, that is clearly abnormal for
a bulb. But a global model trained on all devices might not catch this,
because 10 megabytes is actually LOW compared to what a camera sends.

So I train a SEPARATE anomaly detection model for each of the 8 device types.
Plus one global fallback model for any unknown device type.

For each device, I use an ENSEMBLE of two algorithms:

Algorithm 1 — Isolation Forest:
This algorithm isolates anomalies by randomly partitioning the feature space.
Anomalies are isolated quickly because they are rare and different.
I use 50 trees with a 5% contamination parameter.
Weight in ensemble: 60%.

Algorithm 2 — One-Class SVM:
This draws a boundary around the normal data in feature space.
Anything outside the boundary is anomalous.
I use RBF kernel with nu=0.05.
Weight in ensemble: 40%.

The ENSEMBLE SCORE is:  0.60 × IsolationForest_score + 0.40 × OC-SVM_score

Both algorithms are trained ONLY on normal flows — they never see anomalous
samples during training. This makes them truly unsupervised detectors.

I also apply a power calibration transform with p=0.6 to spread the scores.
This pushes normal traffic scores close to 0, and anomalous traffic scores
above the alert threshold of 0.75.

If the score is above 0.75, the traffic is flagged as anomalous.
The severity is then classified into four levels:
Low (0.50–0.65), Medium (0.65–0.80), High (0.80–0.90), Critical (0.90–1.00).
In practice, since anomalies are only flagged above 0.75, the effective range
starts from Medium severity.
""")

key_points([
    "Per-device anomaly models — 8 separate models + 1 global fallback",
    "Ensemble: 60% Isolation Forest + 40% One-Class SVM",
    "Trained ONLY on normal traffic — unsupervised detection",
    "Alert threshold = 0.75 — above this = anomalous",
    "4 severity levels: Low / Medium / High / Critical",
    "Power calibration (p=0.6) separates normal and anomalous score regions",
])

likely_questions([
    ("Why use Isolation Forest and One-Class SVM together?",
     "Isolation Forest works well in high-dimensional spaces and is fast, "
     "but it can miss some boundary cases. One-Class SVM creates a tighter "
     "boundary around normal data. Their combination gives better precision "
     "and recall than either alone. The 60/40 weighting was determined by "
     "empirical testing on the validation set."),
    ("How did you choose the 0.75 threshold?",
     "I analyzed the score distribution on normal and anomalous flows and "
     "found that normal scores cluster below 0.7 and anomalous scores cluster "
     "above 0.8. The 0.75 threshold gives a good balance between precision "
     "and recall with minimal false positives."),
    ("What if a new device type not in training is connected?",
     "The global fallback model handles this case. It was trained on normal "
     "traffic from ALL 8 device types combined, so it has a broad baseline "
     "of what normal network behaviour looks like."),
])

transition("Anomaly Types Detected")

# ── SLIDE 11 ─────────────────────────────────────────────────────────────────
slide_section(11, "Anomaly Types Detected", "~2 min")

what_to_say("""
Let me explain the three specific attack types my system detects.

ATTACK TYPE 1 — Data Exfiltration:
This is when a device secretly uploads large amounts of data to an outside server.

The signature of this attack in the features is:
— Upload bytes suddenly increases by 30 to 80 times the normal amount
— Byte rate increases by 25 to 60 times normal
— Upload-to-download ratio becomes extremely high — data is going out, not coming in

Example: A Smart Camera normally uploads about 28 megabytes per flow.
During a data exfiltration attack, it might upload 1.5 gigabytes —
that is 50 times the normal amount. My system detects this immediately
and raises a Critical alert.

ATTACK TYPE 2 — Port Scan:
This is when a compromised device starts probing your home network,
looking for other devices with open ports — like a burglar trying every door.

The signature:
— Unique destination ports jumps from 2–3 ports to 500–1,500 ports
— Port entropy becomes very high — close to 9 or 10 out of maximum 10
— IP entropy also increases because many hosts are being probed
— Packet count increases 20 to 50 times

Example: A Smart Bulb normally talks to only 1 port — port 5683 for CoAP.
If it suddenly starts hitting 800 different ports, my system flags it immediately.

ATTACK TYPE 3 — DoS Participation:
This is when a compromised device is being used as part of a botnet
to flood a target server with traffic, making it unavailable.

The signature:
— Packet rate increases by 50 to 120 times
— Packet count increases by 40 to 100 times
— SYN ratio becomes very high — this is the classic SYN flood attack signature

Example: A Motion Sensor normally sends 10 packets per second.
During DoS participation, it might send 150,000 packets per second.
That is 15,000 times the normal rate. My system catches this instantly.
""")

key_points([
    "Data Exfiltration: 30–80× upload spike, high byte rate",
    "Port Scan: 500–1500 unique destination ports, high port entropy (8–10)",
    "DoS Participation: 50–120× packet rate, high SYN ratio",
    "All anomalies produce score > 0.75 → alert is raised",
])

likely_questions([
    ("Can these attacks be missed if the anomaly is gradual?",
     "Gradual attacks are harder to detect with threshold-based methods. "
     "However, the Isolation Forest component handles this better because "
     "it detects statistical outliers relative to the training distribution, "
     "even for smaller deviations. A gradual exfiltration might score around "
     "0.78 — which still triggers a Low severity alert."),
])

transition("REST API and Dashboard")

# ── SLIDE 12 ─────────────────────────────────────────────────────────────────
slide_section(12, "REST API & Live Dashboard", "~2 min")

what_to_say("""
My framework is not just a research prototype — it is a fully deployable system
with a live REST API and a real-time monitoring dashboard.

I built a complete REST API using FastAPI, running on Port 8000.
FastAPI is a modern Python framework that automatically generates API documentation
and is very fast because it is asynchronous.

The API has 9 endpoints:
— GET /health: checks if the models are loaded and the system is running
— GET /devices: lists all 8 supported device types with descriptions
— POST /fingerprint: takes 37 features and returns the device type and confidence score
— POST /anomaly/score: takes features and a device type, returns the anomaly score
— POST /analyze: combines both — one call gives you device type AND anomaly score
— GET /alerts/recent: returns the last 50 alerts with severity and device info
— GET /metrics: returns uptime, request count, anomaly rate, and SHAP status
— POST /demo/inject: injects a simulated anomaly — useful for live demonstration
— POST /explain: NEW — takes 37 features and returns SHAP values for the top 10 features
  that drove the prediction. This makes every single prediction fully explainable.

I also built a custom dark-themed Swagger UI at the /docs endpoint.
This gives a beautiful cybersecurity-styled API documentation page with
a stats bar showing Models Loaded, 8 Devices, 37 Features, 100% Accuracy, and SHAP Ready.

For the live dashboard, I used Plotly Dash on Port 8050.
It shows real-time charts for anomaly scores, device breakdowns,
alert severity histograms, a live alert table, and — most importantly —
a SHAP Explainability panel at the bottom.

This SHAP panel calls POST /explain automatically every 5 seconds.
It shows a horizontal bar chart: green bars mean a feature is PUSHING the model
toward the predicted device, red bars mean it is pushing AWAY.
So when you watch the dashboard, you can see in real time exactly WHY
the model classified the current device as, say, a smart camera.
This is what we call Explainable AI — the model is not a black box.
""")

key_points([
    "FastAPI on Port 8000 — 9 endpoints, async, auto-documented",
    "POST /explain = live SHAP endpoint — returns top 10 feature contributions",
    "Custom dark Swagger UI at /docs — JetBrains Mono, cyan/purple theme",
    "Plotly Dash dashboard on Port 8050 — auto-refreshes every 5 seconds",
    "SHAP panel in dashboard shows green/red bars explaining every prediction live",
    "POST /analyze = one call for fingerprinting + anomaly detection together",
])

transition("Results & Performance")

# ── SLIDE 13 ─────────────────────────────────────────────────────────────────
slide_section(13, "Results & Performance Metrics", "~2 min")

what_to_say("""
Now let me present the results.

For DEVICE FINGERPRINTING:
My Random Forest model achieved 100% accuracy on the test set.
The ROC-AUC score is 1.000 — perfect score.
All 8 device types were classified with 100% precision, recall, and F1-score.

This means: given the 37 features of any network flow, my system can
correctly identify which device type it belongs to — with zero errors.

The confidence threshold of 0.75 means the system only commits to a prediction
when it is at least 75% confident. Below that, it returns "unknown" — which is
safer than a wrong prediction in a security context.

For ANOMALY DETECTION:
The per-device IF + OC-SVM ensemble achieves high ROC-AUC across all device types.
The values range from around 0.91 to 0.97.

This means:
— The system correctly identifies anomalous traffic in most cases
— False positives — normal traffic flagged as anomalous — are very low
— Each device type has its own model tuned to its specific normal behaviour

The average API response time is under 50 milliseconds — fast enough for
real-time monitoring.

The system processes one prediction per API request. In a production deployment,
this can be parallelized to handle thousands of flows per second.

In summary: The framework successfully achieves all 6 objectives I set out with.
Device fingerprinting is perfect, and anomaly detection is highly accurate.
""")

key_points([
    "Fingerprinting: 100% accuracy, ROC-AUC = 1.000 on all 8 device types",
    "Anomaly detection ROC-AUC: 0.91 to 0.97 per device type",
    "API response: < 50 ms per prediction",
    "Confidence threshold 0.75 prevents wrong predictions — returns 'unknown'",
])

likely_questions([
    ("How would accuracy change with real noisy traffic?",
     "Real traffic has noise — retransmissions, background processes, variable loads. "
     "I expect accuracy to drop to around 92–96% in practice. The device signatures "
     "are still highly distinct, so the model should generalize well. Testing on real "
     "traffic capture is part of my future work."),
    ("What is the false positive rate for anomaly detection?",
     "On the synthetic test set, the false positive rate is very low — below 2%. "
     "Normal traffic scores cluster below 0.70, well below the 0.75 threshold. "
     "The power calibration step specifically addresses this by separating the score "
     "distributions of normal and anomalous traffic."),
])

transition("Technology Stack")

# ── SLIDE 14 ─────────────────────────────────────────────────────────────────
slide_section(14, "Technology Stack", "~1 min")

what_to_say("""
Let me briefly mention the technologies I used.

The entire framework is built in Python 3.9.
For machine learning, I used scikit-learn — which provides Random Forest,
Gradient Boosting, SVM, Isolation Forest, and One-Class SVM.

For class balancing, I used imbalanced-learn with SMOTE.

For the API layer, I used FastAPI with Uvicorn — an asynchronous web server.

For the dashboard, I used Plotly Dash — a Python framework for building
interactive web applications with charts.

For data processing, I used NumPy and Pandas.
For logging, I used Loguru — a structured logging library.
For saving and loading models, I used joblib.
For visualization in training reports, I used Matplotlib and Seaborn.
For testing, I wrote a test suite using pytest.

The entire system runs on a Windows 11 machine with 8GB RAM —
no GPU required, no cloud services, completely local.
""")

key_points([
    "Python 3.9 — entire project",
    "scikit-learn: RF, GB, SVM, IsolationForest, OC-SVM",
    "FastAPI + Uvicorn: REST API",
    "Plotly Dash: live dashboard",
    "joblib: model save/load | loguru: logging | pytest: testing",
    "Runs on 8GB RAM Windows 11 — no GPU needed",
])

transition("Conclusion & Future Work")

# ── SLIDE 15 ─────────────────────────────────────────────────────────────────
slide_section(15, "Conclusion & Future Work", "~2 min")

what_to_say("""
Let me conclude.

WHAT I ACHIEVED:

I successfully designed and implemented a complete ML-based framework for
IoT device fingerprinting and anomaly detection.

I extracted 37 carefully crafted features that capture the full behavioral
signature of IoT devices without looking inside their packets.

My Random Forest classifier achieves 100% accuracy in identifying 8 different
smart home device types.

My per-device Isolation Forest + One-Class SVM ensemble successfully detects
three types of attacks — data exfiltration, port scanning, and DoS participation.

The framework is deployed as a production-ready REST API with a live dashboard —
it is not just a research prototype, it can actually be used.

WHAT I PLAN TO DO NEXT:

First, I want to capture REAL traffic from physical IoT devices using Wireshark
and validate my model on actual device traffic.

Second, I want to explore Deep Learning — specifically LSTM networks for
temporal anomaly detection, and Autoencoders for unsupervised learning.

Third, I am interested in Federated Learning — where multiple homes each
train their own model locally, and the models are combined without sharing
private data. This is important for user privacy.

Fourth, I want to deploy this on a Raspberry Pi at the home gateway,
so it runs locally at the router level without any cloud dependency.

Thank you.
""")

key_points([
    "All 6 objectives achieved — fingerprinting 100%, anomaly detection high AUC",
    "Complete deployable system: API + Dashboard",
    "Future: real traffic, LSTM/Autoencoder, Federated Learning, edge deployment",
])

# ── SLIDE 16 ─────────────────────────────────────────────────────────────────
slide_section(16, "Thank You / Q&A", "~5-10 min Q&A")

what_to_say("""
Thank you very much for your patient attention.

To summarize in one sentence:

I built a system that watches smart home network traffic, automatically identifies
what device is generating it, and immediately alerts you if that device starts
behaving in a suspicious or dangerous way.

I am now open for questions. Please feel free to ask anything about
my methodology, models, results, or future plans.

[ Smile. Stand straight. Wait for questions calmly. ]
""")

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("QUICK REFERENCE — NUMBERS TO REMEMBER")
set_font(r, size=13, bold=True, color=(0, 180, 200))
shade_paragraph(p, "E8F8FF")

numbers = [
    ("1,600",   "Total network flows in dataset"),
    ("8",       "IoT device types supported"),
    ("200",     "Flows per device"),
    ("37",      "Features extracted per flow"),
    ("7",       "Feature categories"),
    ("5%",      "Anomaly fraction in dataset"),
    ("100",     "Number of trees in Random Forest"),
    ("15",      "Max depth of Random Forest"),
    ("0.75",    "Confidence threshold (fingerprinting) AND alert threshold (anomaly)"),
    ("50",      "Number of trees in Isolation Forest"),
    ("0.05",    "Contamination parameter (IF) and nu (OC-SVM)"),
    ("60/40",   "Ensemble weighting: IF 60%, OC-SVM 40%"),
    ("100%",    "Fingerprinting accuracy on test set"),
    ("1.000",   "ROC-AUC for fingerprinting"),
    ("0.91–0.97","Anomaly detection ROC-AUC range across device types"),
    ("<50ms",   "API response time per prediction"),
    ("8000",    "FastAPI port"),
    ("8050",    "Plotly Dash dashboard port"),
    ("4",                "Alert severity levels: Low(0.50–0.65) / Med(0.65–0.80) / High(0.80–0.90) / Critical(0.90+)"),
    ("3",                "Anomaly types: Exfiltration / Port Scan / DoS"),
    ("70/15/15",         "Train / Validation / Test split"),
    ("RobustScaler",     "Scaler used — NOT StandardScaler (outlier-resistant, uses median+IQR)"),
]

table = doc.add_table(rows=len(numbers)+1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Value"
hdr[1].text = "What It Means"
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            set_font(r, size=11, bold=True, color=(0, 100, 140))

for i, (val, meaning) in enumerate(numbers):
    row = table.rows[i+1].cells
    row[0].text = val
    row[1].text = meaning
    for j, cell in enumerate(row):
        for para in cell.paragraphs:
            for r in para.runs:
                color = (0, 160, 180) if j == 0 else (50, 50, 50)
                set_font(r, size=11, bold=(j == 0), color=color)

doc.add_paragraph()

# Final tips
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("FINAL PRESENTATION TIPS")
set_font(r, size=12, bold=True, color=(180, 80, 0))
shade_paragraph(p, "FFF5EE")

final_tips = [
    "Start strong — say your name, roll number, and title clearly in the first 20 seconds.",
    "Speak slowly. Pause after important sentences. The panel needs time to absorb.",
    "When showing a chart or table, point to specific numbers — don't just say 'as you can see'.",
    "If asked something you don't know, say: 'That is an excellent point. I have not explored that yet but it would be a great extension.'",
    "Never argue with a panelist. If they challenge your choice, say: 'That is a valid concern. My reason for this choice was...'",
    "Stay within time — 20 minutes presentation, 5-10 minutes Q&A.",
    "Keep a glass of water nearby. It is okay to take a sip before answering a tough question.",
    "End with confidence — 'I am happy to take any further questions.'",
]
for t in final_tips:
    p2 = doc.add_paragraph(style="List Bullet")
    r2 = p2.add_run(t)
    set_font(r2, size=11)

# Save
out = r"D:\IoT_Device_Fingerprinting_Framework\Presentation_Script.docx"
doc.save(out)
print(f"Saved: {out}")
