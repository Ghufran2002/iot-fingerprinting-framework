"""
Professional fixes for DISSERTATION_REPORT.docx:
1. Fix Section 8.3.1 style → Heading 3 (currently Normal)
2. Fix SHAP body paragraphs style → proper Body Text
3. Add POST /explain to Appendix B with full JSON example
4. Fill LIST OF ABBREVIATIONS (was empty)
5. Add SHAP panel to Chapter 7.5 Dashboard
6. Add SHAP to Chapter 9.1 Summary of Contributions
7. Fix dashboard refresh "3 seconds" → "5 seconds"
8. Fix Chapter 7.4 note style → Body Text
Run: python fix_report_professional.py
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document("DISSERTATION_REPORT.docx")
body = doc.element.body
all_p = doc.paragraphs   # live list


# ── Helper: insert paragraph with style after a given paragraph object ────────
def insert_after(anchor_para, text, style_name="Body Text"):
    idx = list(body).index(anchor_para._element)
    new_p = doc.add_paragraph(text, style=style_name)
    body.remove(new_p._element)
    body.insert(idx + 1, new_p._element)
    return new_p


def insert_after_idx(idx_in_body, text, style_name="Body Text"):
    new_p = doc.add_paragraph(text, style=style_name)
    body.remove(new_p._element)
    body.insert(idx_in_body + 1, new_p._element)
    return new_p


def find_para(search_text, exact=False):
    for p in doc.paragraphs:
        if exact:
            if p.text.strip() == search_text:
                return p
        else:
            if search_text in p.text:
                return p
    return None


# ════════════════════════════════════════════════════════════════════════════
# FIX 1: Section 8.3.1 style — Normal → Heading 3
# ════════════════════════════════════════════════════════════════════════════
print("Fix 1: 8.3.1 style...")
p_831 = find_para("8.3.1 SHAP-Based Explainability Analysis")
if p_831:
    p_831.style = doc.styles["Heading 3"]
    print("  Done.")
else:
    print("  NOT FOUND")


# ════════════════════════════════════════════════════════════════════════════
# FIX 2: SHAP body paragraphs in 8.3.1 — Normal → First Paragraph / Body Text
# ════════════════════════════════════════════════════════════════════════════
print("Fix 2: SHAP body paragraph styles...")
shap_anchors = [
    "To address the interpretability requirement of a forensic-grade security system",
    "The SHAP TreeExplainer is applied to the trained Random Forest model",
    "Global Feature Rankings (Mean |SHAP|",
    "Per-Device Analysis (SHAP Heatmap):",
    "Live Explainability Deployment:",
]
first = True
for anchor in shap_anchors:
    p = find_para(anchor)
    if p:
        p.style = doc.styles["First Paragraph"] if first else doc.styles["Body Text"]
        first = False
        print(f"  Fixed: {anchor[:50]}")
    else:
        print(f"  NOT FOUND: {anchor[:50]}")


# ════════════════════════════════════════════════════════════════════════════
# FIX 3: Chapter 7.4 Note style — Normal → Body Text
# ════════════════════════════════════════════════════════════════════════════
print("Fix 3: Chapter 7.4 note style...")
p_note = find_para("Note: Table 7.2 lists all 9 endpoints")
if p_note:
    p_note.style = doc.styles["Body Text"]
    print("  Done.")


# ════════════════════════════════════════════════════════════════════════════
# FIX 4: Chapter 7.5 Dashboard — add SHAP panel + fix refresh rate
# ════════════════════════════════════════════════════════════════════════════
print("Fix 4: Chapter 7.5 Dashboard — SHAP panel + refresh rate...")

# Fix "3 seconds" → "5 seconds"
p_dash = find_para("polls the FastAPI backend via HTTP at a configurable interval (default: 3 seconds)")
if p_dash:
    for run in p_dash.runs:
        if "3 seconds" in run.text:
            run.text = run.text.replace("3 seconds", "5 seconds")
    print("  Refresh rate fixed: 3s → 5s")

# Add SHAP panel entry to dashboard panel list
p_demo = find_para("Demo Controls: Buttons to inject simulated attack traffic")
if p_demo:
    # Insert SHAP panel entry after Demo Controls
    shap_dash = insert_after(
        p_demo,
        "SHAP Explainability Panel: Live horizontal bar chart displaying per-feature SHAP "
        "contributions for the most recent prediction. Green bars indicate features pushing "
        "the model toward the identified device type; red bars indicate features pushing away. "
        "The panel polls POST /explain every 5 seconds, providing real-time Explainable AI "
        "output alongside the anomaly monitoring stream.",
        "Compact"
    )
    print("  SHAP panel added to 7.5.")
else:
    print("  Demo Controls para not found.")


# ════════════════════════════════════════════════════════════════════════════
# FIX 5: Appendix B — Add POST /explain endpoint
# ════════════════════════════════════════════════════════════════════════════
print("Fix 5: Appendix B — adding POST /explain...")

p_metrics = find_para("Returns aggregate metrics including uptime")
if p_metrics:
    entries = [
        ("POST /explain", "Heading 3"),
        (
            "Returns SHAP-based feature attributions for the predicted device type. "
            "Accepts a FlowFeatures payload (37 features). Loads SHAP TreeExplainer on the "
            "Random Forest model at API startup for zero-latency per-request inference.",
            "First Paragraph"
        ),
        ("Request body: FlowFeatures (all 37 features — same schema as POST /fingerprint)", "Body Text"),
        ("Response:", "Body Text"),
        (
            '{\n'
            '  "device_type": "smart_camera",\n'
            '  "confidence": 0.9823,\n'
            '  "explanation": [\n'
            '    {"feature": "tcp_ratio",      "shap_value":  0.238607, "direction": "toward", "abs_impact": 0.238607},\n'
            '    {"feature": "udp_ratio",      "shap_value":  0.224243, "direction": "toward", "abs_impact": 0.224243},\n'
            '    {"feature": "ack_ratio",      "shap_value":  0.094926, "direction": "toward", "abs_impact": 0.094926},\n'
            '    {"feature": "is_https",       "shap_value":  0.031831, "direction": "toward", "abs_impact": 0.031831},\n'
            '    {"feature": "mean_dest_port", "shap_value":  0.020067, "direction": "toward", "abs_impact": 0.020067},\n'
            '    {"feature": "upload_bytes",   "shap_value": -0.026657, "direction": "away",   "abs_impact": 0.026657}\n'
            '  ],\n'
            '  "note": "Positive SHAP = pushes prediction TOWARD this device; negative = pushes AWAY"\n'
            '}',
            "Source Code"
        ),
    ]
    # Insert in reverse so order is maintained
    anchor = p_metrics
    for text, style in reversed(entries):
        insert_after(anchor, text, style)
    print("  POST /explain added to Appendix B.")
else:
    print("  GET /metrics para not found.")


# ════════════════════════════════════════════════════════════════════════════
# FIX 6: LIST OF ABBREVIATIONS — fill it
# ════════════════════════════════════════════════════════════════════════════
print("Fix 6: LIST OF ABBREVIATIONS...")

p_abbr = find_para("LIST OF ABBREVIATIONS", exact=True)
if not p_abbr:
    p_abbr = find_para("LIST OF ABBREVIATIONS")

if p_abbr:
    abbreviations = [
        ("AI",       "Artificial Intelligence"),
        ("API",      "Application Programming Interface"),
        ("AUC",      "Area Under the Curve"),
        ("CoAP",     "Constrained Application Protocol"),
        ("CORS",     "Cross-Origin Resource Sharing"),
        ("DPI",      "Deep Packet Inspection"),
        ("DoS",      "Denial of Service"),
        ("HTTP",     "Hypertext Transfer Protocol"),
        ("HTTPS",    "Hypertext Transfer Protocol Secure"),
        ("IF",       "Isolation Forest"),
        ("IoT",      "Internet of Things"),
        ("IP",       "Internet Protocol"),
        ("JSON",     "JavaScript Object Notation"),
        ("KPI",      "Key Performance Indicator"),
        ("LSTM",     "Long Short-Term Memory"),
        ("ML",       "Machine Learning"),
        ("MQTT",     "Message Queuing Telemetry Transport"),
        ("NTP",      "Network Time Protocol"),
        ("OC-SVM",   "One-Class Support Vector Machine"),
        ("REST",     "Representational State Transfer"),
        ("RF",       "Random Forest"),
        ("ROC",      "Receiver Operating Characteristic"),
        ("RobustScaler", "Robust feature normalisation using interquartile range"),
        ("SHAP",     "SHapley Additive exPlanations"),
        ("SMOTE",    "Synthetic Minority Over-sampling Technique"),
        ("SVM",      "Support Vector Machine"),
        ("TCP",      "Transmission Control Protocol"),
        ("UDP",      "User Datagram Protocol"),
        ("URL",      "Uniform Resource Locator"),
        ("XAI",      "Explainable Artificial Intelligence"),
    ]
    anchor = p_abbr
    for abbr, full in reversed(abbreviations):
        p_new = insert_after(anchor, f"{abbr:<15} {full}", "Compact")
    print(f"  Added {len(abbreviations)} abbreviations.")
else:
    print("  LIST OF ABBREVIATIONS not found.")


# ════════════════════════════════════════════════════════════════════════════
# FIX 7: Chapter 9.1 Summary of Contributions — add SHAP
# ════════════════════════════════════════════════════════════════════════════
print("Fix 7: Chapter 9.1 — SHAP contribution...")

p_contrib = find_para(
    "A multi-algorithm fingerprinting subsystem comparing Random Forest, Gradient Boosting, SVM, and Voti"
)
if not p_contrib:
    p_contrib = find_para("multi-algorithm fingerprinting subsystem")

if p_contrib:
    insert_after(
        p_contrib,
        "A live SHAP (SHapley Additive exPlanations) explainability layer integrated directly into "
        "the REST API via POST /explain. This endpoint provides per-prediction, mathematically "
        "grounded feature attributions satisfying the axioms of efficiency, symmetry, and null "
        "player — transforming the system from a black-box classifier into a fully auditable, "
        "forensic-grade decision support tool. The SHAP panel in the monitoring dashboard "
        "renders these attributions as a live green/red bar chart updating every 5 seconds.",
        "Normal"
    )
    print("  SHAP contribution added to Chapter 9.1.")
else:
    print("  Contribution anchor not found.")


# ════════════════════════════════════════════════════════════════════════════
# FIX 8: /health response in Appendix B — add shap_ready field
# ════════════════════════════════════════════════════════════════════════════
print("Fix 8: /health response — add shap_ready...")
for p in doc.paragraphs:
    if '"models_loaded": true' in p.text and '"uptime_seconds"' in p.text:
        old = p.text
        new = old.replace(
            '"uptime_seconds": 3612.4,\n  "request_count": 847',
            '"shap_ready": true,\n  "uptime_seconds": 3612.4,\n  "request_count": 847'
        )
        if new != old:
            for run in p.runs:
                run.text = ""
            p.runs[0].text = new
            print("  shap_ready added to /health response.")
        break


# ════════════════════════════════════════════════════════════════════════════
doc.save("DISSERTATION_REPORT.docx")
print("\nAll fixes applied. Saved DISSERTATION_REPORT.docx")
