"""Fix severity case and Demo 8 confidence in PANEL_DEMO_GUIDE.docx."""
from docx import Document

path = r"D:\IoT_Device_Fingerprinting_Framework\PANEL_DEMO_GUIDE.docx"
doc = Document(path)

replacements = [
    ('"severity": "High"',     '"severity": "HIGH"'),
    ('"severity": "Critical"', '"severity": "CRITICAL"'),
    ('"severity": "Medium"',   '"severity": "MEDIUM"'),
    ('"severity": "Low"',      '"severity": "LOW"'),
    # Demo 8 confidence
    ('"confidence": 0.9900',   '"confidence": 0.9600'),
    # Demo 8 say-to-panel: 99% → 96% and High-severity → HIGH-severity
    ("smart camera with 99% confidence", "smart camera with 96% confidence"),
    ("raises a High-severity alert",     "raises a HIGH-severity alert"),
]

changes = 0

def fix_para(para):
    global changes
    for old, new in replacements:
        if old in para.text:
            full = para.text
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    changes += 1
                    return
            # cross-run fallback
            if old in full and para.runs:
                para.runs[0].text = full.replace(old, new)
                for r in para.runs[1:]:
                    r.text = ""
                changes += 1

for para in doc.paragraphs:
    fix_para(para)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                fix_para(para)

doc.save(path)
print(f"Done — {changes} replacements in {path}")
